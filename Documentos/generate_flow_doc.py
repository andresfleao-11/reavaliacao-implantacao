"""
Script para gerar documentação DOCX do fluxo de cotação via Google Shopping
"""
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading_style(doc, heading_text, level=1):
    """Adiciona heading com estilo"""
    heading = doc.add_heading(heading_text, level=level)
    return heading

def add_code_block(doc, code_text):
    """Adiciona bloco de código"""
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Cm(1)
    return p

def create_documentation():
    doc = Document()

    # Configurar estilos
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Título Principal
    title = doc.add_heading('Fluxo de Cotação via Google Shopping', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Sistema de Reavaliacao Patrimonial')
    doc.add_paragraph('Documentação Técnica do Fluxo de Busca e Validação de Preços')
    doc.add_paragraph()

    # ========================================
    # 1. VISAO GERAL
    # ========================================
    add_heading_style(doc, '1. Visão Geral do Fluxo', 1)

    doc.add_paragraph(
        'O sistema de cotação realiza buscas de preços em lojas online brasileiras através da API '
        'do Google Shopping (via SerpAPI). O fluxo foi otimizado para minimizar chamadas de API '
        'e garantir preços consistentes através de um sistema de blocos de variação.'
    )

    doc.add_paragraph('Etapas principais:')
    items = [
        '1. Análise da imagem/texto do item (IA: Claude ou GPT)',
        '2. Geração da query de busca otimizada',
        '3. Chamada única ao Google Shopping API',
        '4. Filtragem de fontes bloqueadas e preços inválidos',
        '5. Criação de blocos de variação',
        '6. Tentativa de obter N cotações válidas por bloco',
        '7. Validação de preço via acesso direto ao site',
        '8. Recálculo de blocos em caso de falhas'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    # ========================================
    # 2. CHAMADA GOOGLE SHOPPING
    # ========================================
    add_heading_style(doc, '2. Chamada Google Shopping API', 1)

    add_heading_style(doc, '2.1 Parâmetros da Requisição', 2)

    doc.add_paragraph('A busca utiliza os seguintes parâmetros:')

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Parâmetro'
    hdr_cells[1].text = 'Valor'
    hdr_cells[2].text = 'Descrição'

    params = [
        ('engine', 'google_shopping', 'Motor de busca'),
        ('q', '[query gerada pela IA]', 'Termo de busca'),
        ('gl', 'br', 'Localização geográfica (Brasil)'),
        ('hl', 'pt-br', 'Idioma'),
        ('google_domain', 'google.com.br', 'Domínio do Google'),
        ('location', '[configurável]', 'Localização SerpAPI (ex: São Paulo)'),
        ('num', '100', 'Máximo de produtos por requisição'),
    ]

    for param, valor, desc in params:
        row_cells = table.add_row().cells
        row_cells[0].text = param
        row_cells[1].text = valor
        row_cells[2].text = desc

    doc.add_paragraph()

    add_heading_style(doc, '2.2 Estrutura da Resposta', 2)

    doc.add_paragraph('A API retorna dois arrays principais:')
    doc.add_paragraph('• shopping_results: Produtos do Google Shopping', style='List Bullet')
    doc.add_paragraph('• inline_shopping_results: Produtos inline (sem link Immersive)', style='List Bullet')

    # ========================================
    # 3. VARIÁVEIS EXTRAÍDAS
    # ========================================
    add_heading_style(doc, '3. Variáveis Extraídas dos Produtos', 1)

    doc.add_paragraph('Para cada produto retornado, são extraídas as seguintes variáveis:')

    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Table Grid'
    hdr = table2.rows[0].cells
    hdr[0].text = 'Campo'
    hdr[1].text = 'Tipo'
    hdr[2].text = 'Descrição'

    vars_list = [
        ('title', 'str', 'Nome/título do produto'),
        ('price', 'str', 'Preço formatado (ex: "R$ 1.299,00")'),
        ('extracted_price', 'float', 'Preço numérico extraído'),
        ('source', 'str', 'Nome da loja/fonte (ex: "Kabum")'),
        ('serpapi_immersive_product_api', 'str', 'URL para API Immersive'),
        ('product_link', 'str', 'Link direto do produto'),
        ('link', 'str', 'Link alternativo'),
    ]

    for campo, tipo, desc in vars_list:
        row = table2.add_row().cells
        row[0].text = campo
        row[1].text = tipo
        row[2].text = desc

    # ========================================
    # 4. FILTROS APLICADOS
    # ========================================
    add_heading_style(doc, '4. Filtros Aplicados aos Produtos', 1)

    add_heading_style(doc, '4.1 Filtro de Fontes Bloqueadas', 2)

    doc.add_paragraph(
        'Fontes com proteção anti-bot são filtradas antes de fazer chamadas à API Immersive, '
        'economizando chamadas de API.'
    )

    doc.add_paragraph('Domínios bloqueados:')
    blocked = [
        'Marketplaces: mercadolivre.com.br, amazon.com.br, aliexpress.com, shopee.com.br, shein.com, wish.com, temu.com',
        'Grandes varejistas: carrefour.com.br, casasbahia.com.br, magazineluiza.com.br, americanas.com.br, submarino.com.br',
    ]
    for b in blocked:
        doc.add_paragraph(b, style='List Bullet')

    add_heading_style(doc, '4.2 Filtro de Domínios Estrangeiros', 2)

    doc.add_paragraph(
        'Apenas lojas brasileiras (.com.br) são permitidas, com exceções para fabricantes globais '
        'que vendem no Brasil (Dell, Lenovo, Samsung, HP, LG, Apple, Asus, Acer).'
    )

    add_heading_style(doc, '4.3 Filtro de Preços Inválidos', 2)

    doc.add_paragraph('Produtos são removidos se:')
    doc.add_paragraph('• extracted_price é None ou zero', style='List Bullet')
    doc.add_paragraph('• Preço não pode ser convertido para número', style='List Bullet')

    add_heading_style(doc, '4.4 Filtro de URLs de Listagem', 2)

    doc.add_paragraph('URLs que representam páginas de busca/categoria são rejeitadas:')
    patterns = ['/busca/', '/search/', '?q=', '/categoria/', '/colecao/', 'buscape.com.br', 'zoom.com.br']
    for p in patterns:
        doc.add_paragraph(f'Padrão: {p}', style='List Bullet')

    # ========================================
    # 5. ORDENAÇÃO
    # ========================================
    add_heading_style(doc, '5. Ordenação de Produtos', 1)

    doc.add_paragraph('Após os filtros, os produtos são ordenados por preço (crescente):')
    add_code_block(doc, 'products.sort(key=lambda x: x.extracted_price)')

    doc.add_paragraph(
        'Em seguida, a lista é limitada a MAX_VALID_PRODUCTS (150 produtos) para evitar '
        'processamento excessivo.'
    )

    # ========================================
    # 6. FORMAÇÃO DE BLOCOS
    # ========================================
    add_heading_style(doc, '6. Lógica de Formação de Blocos', 1)

    add_heading_style(doc, '6.1 Conceito de Bloco de Variação', 2)

    doc.add_paragraph(
        'Um bloco de variação é um conjunto de produtos consecutivos (ordenados por preço) '
        'onde o produto mais caro está dentro da variação máxima permitida em relação ao '
        'produto mais barato do bloco.'
    )

    doc.add_paragraph('Fórmula da variação:')
    add_code_block(doc, 'variacao = (preco_maximo / preco_minimo - 1) * 100')

    doc.add_paragraph(
        'Por exemplo, com variacao_maxima = 25%, um bloco iniciando em R$ 100 pode incluir '
        'produtos até R$ 125.'
    )

    add_heading_style(doc, '6.2 Algoritmo de Criação (Sliding Window)', 2)

    doc.add_paragraph('O sistema usa abordagem de janela deslizante:')

    steps = [
        '1. Para cada produto na lista ordenada, criar um bloco iniciando nele',
        '2. Calcular preço máximo permitido: min_price * (1 + variacao_maxima)',
        '3. Incluir todos produtos consecutivos até que um exceda o máximo',
        '4. Descartar blocos com menos produtos que o N configurado'
    ]
    for s in steps:
        doc.add_paragraph(s, style='List Number')

    add_heading_style(doc, '6.3 Exemplo Prático', 2)

    doc.add_paragraph('Produtos ordenados: [R$ 100, R$ 102, R$ 104, R$ 110, R$ 125, R$ 130, R$ 140]')
    doc.add_paragraph('Configuração: N=3 cotações, variação máxima=25%')
    doc.add_paragraph()

    table3 = doc.add_table(rows=1, cols=4)
    table3.style = 'Table Grid'
    hdr3 = table3.rows[0].cells
    hdr3[0].text = 'Bloco'
    hdr3[1].text = 'Produtos'
    hdr3[2].text = 'Faixa'
    hdr3[3].text = 'Válido?'

    blocos = [
        ('1', '[100, 102, 104, 110, 125]', 'R$ 100 - R$ 125', 'Sim (5 >= 3)'),
        ('2', '[102, 104, 110, 125]', 'R$ 102 - R$ 125', 'Sim (4 >= 3)'),
        ('3', '[104, 110, 125, 130]', 'R$ 104 - R$ 130', 'Sim (4 >= 3)'),
        ('4', '[110, 125, 130]', 'R$ 110 - R$ 137.5', 'Sim (3 >= 3)'),
        ('5', '[125, 130, 140]', 'R$ 125 - R$ 156', 'Sim (3 >= 3)'),
    ]

    for b, prods, faixa, valido in blocos:
        row = table3.add_row().cells
        row[0].text = b
        row[1].text = prods
        row[2].text = faixa
        row[3].text = valido

    add_heading_style(doc, '6.4 Priorização de Blocos', 2)

    doc.add_paragraph('Blocos são ordenados por:')
    doc.add_paragraph('1. Maior quantidade de produtos (mais chances de sucesso)', style='List Number')
    doc.add_paragraph('2. Menor preço inicial (opções mais baratas primeiro)', style='List Number')

    add_code_block(doc, 'sorted_blocks = sorted(blocks, key=lambda b: (-len(b), b[0].extracted_price))')

    # ========================================
    # 7. CONCEITO DE FALHA DE PRODUTO
    # ========================================
    add_heading_style(doc, '7. Conceito de Falha de Produto', 1)

    doc.add_paragraph(
        'Um produto é marcado como "falha" quando não pode ser usado como cotação válida. '
        'Produtos que falham são removidos da lista para recálculo dos blocos.'
    )

    add_heading_style(doc, '7.1 Causas de Falha de Produto', 2)

    table4 = doc.add_table(rows=1, cols=3)
    table4.style = 'Table Grid'
    hdr4 = table4.rows[0].cells
    hdr4[0].text = 'Tipo de Falha'
    hdr4[1].text = 'Código'
    hdr4[2].text = 'Descrição'

    falhas = [
        ('Sem link de loja', 'no_store_link', 'API Immersive não retornou link válido'),
        ('Domínio bloqueado', 'blocked_domain', 'Loja está na lista de bloqueio'),
        ('Domínio estrangeiro', 'foreign_domain', 'Loja não é brasileira'),
        ('Domínio duplicado', 'duplicate_domain', 'Já existe cotação desta loja'),
        ('URL de listagem', 'listing_url', 'URL é página de busca, não produto'),
        ('Divergência de preço', 'price_mismatch', 'Preço do site difere > 15% do Google'),
        ('Erro de extração', 'extraction_error', 'Não foi possível extrair preço do site'),
    ]

    for tipo, codigo, desc in falhas:
        row = table4.add_row().cells
        row[0].text = tipo
        row[1].text = codigo
        row[2].text = desc

    add_heading_style(doc, '7.2 Validação de Preço (Price Mismatch)', 2)

    doc.add_paragraph(
        'Quando o sistema acessa o site da loja para capturar screenshot e extrair preço, '
        'ele compara o preço encontrado com o preço informado pelo Google Shopping.'
    )

    doc.add_paragraph('Regra de validação:')
    add_code_block(doc, '''price_diff_percent = abs(preco_site - preco_google) / preco_google * 100
if price_diff_percent > 15:
    # Divergência muito alta - usar preço do Google
    # Marca produto para validação manual se necessário''')

    doc.add_paragraph(
        'Se a diferença for maior que 15%, o sistema pode usar o preço do Google como fallback '
        'ou marcar o produto como falha, dependendo do contexto.'
    )

    # ========================================
    # 8. CONCEITO DE FALHA DE BLOCO
    # ========================================
    add_heading_style(doc, '8. Conceito de Falha de Bloco', 1)

    doc.add_paragraph(
        'Um bloco é considerado "falho" quando não consegue fornecer o número mínimo de '
        'cotações válidas (N configurado).'
    )

    add_heading_style(doc, '8.1 Causas de Falha de Bloco', 2)

    doc.add_paragraph('• Todos os produtos do bloco falharam individualmente', style='List Bullet')
    doc.add_paragraph('• Produtos restantes após falhas são insuficientes (< N)', style='List Bullet')
    doc.add_paragraph('• Não há produtos não-tentados suficientes no bloco', style='List Bullet')

    add_heading_style(doc, '8.2 Comportamento Após Falha de Bloco', 2)

    doc.add_paragraph('Quando um bloco falha:')
    doc.add_paragraph('1. Os produtos válidos obtidos são preservados', style='List Number')
    doc.add_paragraph('2. Os produtos que falharam são removidos da lista', style='List Number')
    doc.add_paragraph('3. Novos blocos são recalculados com produtos restantes', style='List Number')
    doc.add_paragraph('4. O próximo bloco prioritário é selecionado', style='List Number')

    # ========================================
    # 9. RECÁLCULO DE BLOCOS
    # ========================================
    add_heading_style(doc, '9. Algoritmo de Recálculo de Blocos', 1)

    doc.add_paragraph(
        'O sistema usa um algoritmo iterativo com recálculo dinâmico de blocos após cada '
        'falha de produto.'
    )

    add_heading_style(doc, '9.1 Fluxo do Algoritmo', 2)

    steps2 = [
        'Iniciar com lista completa de produtos filtrados',
        'Criar blocos de variação (sliding window)',
        'Selecionar bloco prioritário',
        'Para cada produto do bloco:',
        '   a. Se já validado anteriormente: reutilizar',
        '   b. Se falhou anteriormente: pular',
        '   c. Senão: tentar obter preço via Immersive API + acesso ao site',
        'Se obteve N cotações: SUCESSO, finalizar',
        'Se houve novas falhas: remover produtos falhos, recalcular blocos',
        'Se não há progresso: verificar sistema de reserva ou encerrar',
        'Repetir até sucesso ou esgotar iterações (max: 15)',
    ]
    for i, s in enumerate(steps2, 1):
        if s.startswith('   '):
            doc.add_paragraph(s.strip(), style='List Bullet')
        else:
            doc.add_paragraph(f'{i}. {s}', style='List Number')

    add_heading_style(doc, '9.2 Sistema de Reserva', 2)

    doc.add_paragraph(
        'Quando um bloco com produtos válidos não tem produtos suficientes não-tentados '
        'para atingir N, o sistema implementa um mecanismo de reserva:'
    )

    doc.add_paragraph('1. Salva os produtos válidos atuais como "reserva"', style='List Number')
    doc.add_paragraph('2. Limpa resultados e tenta um bloco alternativo sem os válidos', style='List Number')
    doc.add_paragraph('3. Se bloco alternativo falha: restaura reserva', style='List Number')
    doc.add_paragraph('4. Continua processamento com resultados restaurados', style='List Number')

    add_heading_style(doc, '9.3 Categorização de Blocos', 2)

    doc.add_paragraph('Durante cada iteração, os blocos são categorizados:')

    table5 = doc.add_table(rows=1, cols=2)
    table5.style = 'Table Grid'
    hdr5 = table5.rows[0].cells
    hdr5[0].text = 'Categoria'
    hdr5[1].text = 'Descrição'

    cats = [
        ('Com todos válidos + suficientes', 'Contém todos produtos válidos anteriores E tem produtos não-tentados suficientes para atingir N'),
        ('Com todos válidos, insuficientes', 'Contém todos válidos mas não tem não-tentados suficientes'),
        ('Sem válidos mas grande', 'Não contém os válidos anteriores mas tem >= N produtos'),
    ]

    for cat, desc in cats:
        row = table5.add_row().cells
        row[0].text = cat
        row[1].text = desc

    doc.add_paragraph()
    doc.add_paragraph('Ordem de prioridade: categoria 1 > categoria 2 > categoria 3')

    # ========================================
    # 10. API IMMERSIVE
    # ========================================
    add_heading_style(doc, '10. Google Immersive Product API', 1)

    doc.add_paragraph(
        'Para obter o link direto da loja (não o link do Google), é necessário chamar a '
        'API Immersive para cada produto.'
    )

    add_heading_style(doc, '10.1 Dados Retornados', 2)

    doc.add_paragraph('A API Immersive retorna:')
    doc.add_paragraph('• product_results.stores[]: Array de lojas com preços', style='List Bullet')
    doc.add_paragraph('• online_sellers[]: Vendedores online alternativos', style='List Bullet')
    doc.add_paragraph('• product_results.link: Link direto do produto (fallback)', style='List Bullet')

    add_heading_style(doc, '10.2 Validação de Preço na API Immersive', 2)

    doc.add_paragraph(
        'Antes de aceitar uma loja da API Immersive, o preço dela é comparado com o '
        'preço do produto original:'
    )

    add_code_block(doc, '''for store in stores:
    store_price = store.extracted_price
    block_price = product.extracted_price

    diff_percent = abs(store_price - block_price) / block_price * 100
    if diff_percent > 15:
        # Pular esta loja - preço muito diferente
        continue

    # Aceitar loja''')

    # ========================================
    # 11. LIMPEZA DE URLs
    # ========================================
    add_heading_style(doc, '11. Limpeza de URLs (Tracking Parameters)', 1)

    doc.add_paragraph(
        'URLs retornadas podem conter parâmetros de rastreamento que causam redirecionamentos '
        'indesejados. O sistema remove estes parâmetros:'
    )

    params_removed = [
        'srsltid (Google Shopping tracking)',
        'pf, mc (tracking genérico)',
        'utm_source, utm_medium, utm_campaign, utm_term, utm_content',
        'gclid, fbclid, ref, ref_',
        '_ga, _gl, dclid',
    ]
    for p in params_removed:
        doc.add_paragraph(p, style='List Bullet')

    # ========================================
    # 12. CONFIGURAÇÕES
    # ========================================
    add_heading_style(doc, '12. Parâmetros Configuráveis', 1)

    table6 = doc.add_table(rows=1, cols=3)
    table6.style = 'Table Grid'
    hdr6 = table6.rows[0].cells
    hdr6[0].text = 'Parâmetro'
    hdr6[1].text = 'Default'
    hdr6[2].text = 'Descrição'

    configs = [
        ('numero_cotacoes_por_pesquisa', '3', 'Quantidade de cotações (N) a obter'),
        ('variacao_maxima_percent', '25', 'Variação máxima permitida (%)'),
        ('serpapi_location', 'São Paulo,...', 'Localização para busca'),
        ('MAX_VALID_PRODUCTS', '150', 'Máximo de produtos a processar'),
        ('MAX_RETRIES', '3', 'Tentativas em caso de rate limit'),
        ('max_iterations', '15', 'Máximo de iterações de recálculo'),
    ]

    for param, default, desc in configs:
        row = table6.add_row().cells
        row[0].text = param
        row[1].text = default
        row[2].text = desc

    # ========================================
    # 13. LOGS E AUDITORIA
    # ========================================
    add_heading_style(doc, '13. Logs e Auditoria', 1)

    doc.add_paragraph(
        'O sistema mantém logs detalhados de cada busca através da estrutura SearchLog:'
    )

    log_fields = [
        ('query', 'Query de busca utilizada'),
        ('limit', 'N configurado'),
        ('variacao_maxima', 'Variação máxima configurada'),
        ('total_raw_products', 'Total de produtos retornados pelo Google'),
        ('after_source_filter', 'Produtos após filtro de fonte'),
        ('blocked_sources', 'Fontes bloqueadas'),
        ('after_price_filter', 'Produtos após filtro de preço'),
        ('invalid_prices', 'Preços inválidos removidos'),
        ('total_blocks_created', 'Total de blocos criados'),
        ('valid_blocks', 'Blocos válidos (>= N produtos)'),
        ('blocks_tried', 'Blocos tentados'),
        ('successful_block_index', 'Índice do bloco que teve sucesso'),
        ('immersive_api_calls', 'Chamadas à API Immersive'),
        ('results_obtained', 'Resultados obtidos'),
        ('block_details[]', 'Detalhes de cada bloco tentado'),
    ]

    for field, desc in log_fields:
        doc.add_paragraph(f'• {field}: {desc}', style='List Bullet')

    # ========================================
    # FIM
    # ========================================
    doc.add_paragraph()
    doc.add_paragraph('---')
    doc.add_paragraph('Documento gerado automaticamente pelo Sistema de Reavaliação Patrimonial')

    return doc

if __name__ == '__main__':
    doc = create_documentation()
    output_path = '/app/Documentos/Fluxo_Google_Shopping.docx'
    doc.save(output_path)
    print(f'Documento gerado: {output_path}')
