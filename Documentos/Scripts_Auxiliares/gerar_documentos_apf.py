"""
Script para gerar documentos de Análise de Pontos de Função
- Documento Word (.docx) com visão executiva
- Planilha Excel (.xlsx) com detalhamento completo
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl
from openpyxl.styles import Font, Fill, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl.chart import BarChart, PieChart, Reference
from datetime import datetime

# ============================================================
# DADOS DA CONTAGEM DE PONTOS DE FUNÇÃO
# ============================================================

# Arquivos Lógicos Internos (ALI)
ALI_DATA = [
    {"id": 1, "nome": "Client (Clientes)", "der": 14, "rlr": 1, "complexidade": "Baixa", "pf": 7, "modulo": "Clientes"},
    {"id": 2, "nome": "Project (Projetos)", "der": 16, "rlr": 2, "complexidade": "Média", "pf": 10, "modulo": "Projetos"},
    {"id": 3, "nome": "QuoteRequest (Cotações)", "der": 18, "rlr": 4, "complexidade": "Alta", "pf": 15, "modulo": "Cotações"},
    {"id": 4, "nome": "QuoteSource (Fontes de Preço)", "der": 11, "rlr": 2, "complexidade": "Baixa", "pf": 7, "modulo": "Cotações"},
    {"id": 5, "nome": "Material (Materiais)", "der": 12, "rlr": 2, "complexidade": "Média", "pf": 10, "modulo": "Materiais"},
    {"id": 6, "nome": "MaterialCharacteristic", "der": 6, "rlr": 1, "complexidade": "Baixa", "pf": 7, "modulo": "Materiais"},
    {"id": 7, "nome": "CharacteristicType", "der": 8, "rlr": 1, "complexidade": "Baixa", "pf": 7, "modulo": "Materiais"},
    {"id": 8, "nome": "Item (Itens Patrimônio)", "der": 11, "rlr": 2, "complexidade": "Média", "pf": 10, "modulo": "Itens"},
    {"id": 9, "nome": "ItemCharacteristic", "der": 5, "rlr": 2, "complexidade": "Baixa", "pf": 7, "modulo": "Itens"},
    {"id": 10, "nome": "User (Usuários)", "der": 7, "rlr": 1, "complexidade": "Baixa", "pf": 7, "modulo": "Usuários"},
    {"id": 11, "nome": "ProjectConfigVersion", "der": 22, "rlr": 3, "complexidade": "Alta", "pf": 15, "modulo": "Configurações"},
    {"id": 12, "nome": "ProjectBankPrice", "der": 7, "rlr": 1, "complexidade": "Baixa", "pf": 7, "modulo": "Configurações"},
    {"id": 13, "nome": "FinancialTransaction", "der": 12, "rlr": 2, "complexidade": "Média", "pf": 10, "modulo": "Financeiro"},
    {"id": 14, "nome": "ApiCostConfig", "der": 11, "rlr": 1, "complexidade": "Baixa", "pf": 7, "modulo": "Financeiro"},
    {"id": 15, "nome": "BankPrice (Banco Preços)", "der": 6, "rlr": 1, "complexidade": "Baixa", "pf": 7, "modulo": "Configurações"},
    {"id": 16, "nome": "RevaluationParam", "der": 5, "rlr": 1, "complexidade": "Baixa", "pf": 7, "modulo": "Configurações"},
    {"id": 17, "nome": "Setting (Configurações)", "der": 3, "rlr": 1, "complexidade": "Baixa", "pf": 7, "modulo": "Configurações"},
    {"id": 18, "nome": "IntegrationSetting", "der": 4, "rlr": 1, "complexidade": "Baixa", "pf": 7, "modulo": "Integrações"},
    {"id": 19, "nome": "File (Arquivos)", "der": 6, "rlr": 1, "complexidade": "Baixa", "pf": 7, "modulo": "Arquivos"},
    {"id": 20, "nome": "GeneratedDocument", "der": 4, "rlr": 2, "complexidade": "Baixa", "pf": 7, "modulo": "Arquivos"},
]

# Arquivos de Interface Externa (AIE)
AIE_DATA = [
    {"id": 1, "nome": "API Anthropic (Claude AI)", "der": 8, "rlr": 1, "complexidade": "Média", "pf": 7, "modulo": "Integrações"},
    {"id": 2, "nome": "API SerpAPI (Google Shopping)", "der": 6, "rlr": 1, "complexidade": "Baixa", "pf": 5, "modulo": "Integrações"},
    {"id": 3, "nome": "Páginas Web (Scraping)", "der": 10, "rlr": 1, "complexidade": "Média", "pf": 7, "modulo": "Integrações"},
]

# Entradas Externas (EE) - Backend
EE_BACKEND = [
    {"id": 1, "nome": "POST /api/clients - Incluir Cliente", "ali": 1, "der": 14, "complexidade": "Média", "pf": 4, "modulo": "Clientes", "camada": "Backend"},
    {"id": 2, "nome": "PUT /api/clients/{id} - Alterar Cliente", "ali": 1, "der": 14, "complexidade": "Média", "pf": 4, "modulo": "Clientes", "camada": "Backend"},
    {"id": 3, "nome": "DELETE /api/clients/{id} - Excluir Cliente", "ali": 1, "der": 2, "complexidade": "Baixa", "pf": 3, "modulo": "Clientes", "camada": "Backend"},
    {"id": 4, "nome": "POST /api/projects - Incluir Projeto", "ali": 2, "der": 16, "complexidade": "Média", "pf": 4, "modulo": "Projetos", "camada": "Backend"},
    {"id": 5, "nome": "PUT /api/projects/{id} - Alterar Projeto", "ali": 2, "der": 16, "complexidade": "Média", "pf": 4, "modulo": "Projetos", "camada": "Backend"},
    {"id": 6, "nome": "DELETE /api/projects/{id} - Excluir Projeto", "ali": 1, "der": 2, "complexidade": "Baixa", "pf": 3, "modulo": "Projetos", "camada": "Backend"},
    {"id": 7, "nome": "POST /api/quotes - Criar Cotação", "ali": 3, "der": 12, "complexidade": "Alta", "pf": 6, "modulo": "Cotações", "camada": "Backend"},
    {"id": 8, "nome": "POST /api/quotes/{id}/cancel - Cancelar Cotação", "ali": 1, "der": 3, "complexidade": "Baixa", "pf": 3, "modulo": "Cotações", "camada": "Backend"},
    {"id": 9, "nome": "POST /api/quotes/{id}/requote - Recotar Item", "ali": 3, "der": 10, "complexidade": "Alta", "pf": 6, "modulo": "Cotações", "camada": "Backend"},
    {"id": 10, "nome": "POST /api/materials - Incluir Material", "ali": 2, "der": 12, "complexidade": "Média", "pf": 4, "modulo": "Materiais", "camada": "Backend"},
    {"id": 11, "nome": "PUT /api/materials/{id} - Alterar Material", "ali": 2, "der": 12, "complexidade": "Média", "pf": 4, "modulo": "Materiais", "camada": "Backend"},
    {"id": 12, "nome": "DELETE /api/materials/{id} - Excluir Material", "ali": 2, "der": 2, "complexidade": "Baixa", "pf": 3, "modulo": "Materiais", "camada": "Backend"},
    {"id": 13, "nome": "POST /api/materials/{id}/characteristics - Incluir Característica", "ali": 2, "der": 6, "complexidade": "Baixa", "pf": 3, "modulo": "Materiais", "camada": "Backend"},
    {"id": 14, "nome": "PUT /api/materials/{id}/characteristics/{id} - Alterar Característica", "ali": 2, "der": 6, "complexidade": "Baixa", "pf": 3, "modulo": "Materiais", "camada": "Backend"},
    {"id": 15, "nome": "DELETE /api/materials/{id}/characteristics/{id} - Excluir Característica", "ali": 1, "der": 2, "complexidade": "Baixa", "pf": 3, "modulo": "Materiais", "camada": "Backend"},
    {"id": 16, "nome": "POST /api/materials/items - Incluir Item", "ali": 3, "der": 11, "complexidade": "Média", "pf": 4, "modulo": "Itens", "camada": "Backend"},
    {"id": 17, "nome": "PUT /api/materials/items/{id} - Alterar Item", "ali": 3, "der": 11, "complexidade": "Média", "pf": 4, "modulo": "Itens", "camada": "Backend"},
    {"id": 18, "nome": "DELETE /api/materials/items/{id} - Excluir Item", "ali": 2, "der": 2, "complexidade": "Baixa", "pf": 3, "modulo": "Itens", "camada": "Backend"},
    {"id": 19, "nome": "POST /api/materials/items/bulk-create - Gerar Itens Lote", "ali": 4, "der": 15, "complexidade": "Alta", "pf": 6, "modulo": "Itens", "camada": "Backend"},
    {"id": 20, "nome": "POST /api/materials/import - Importar CSV/XLSX", "ali": 4, "der": 12, "complexidade": "Alta", "pf": 6, "modulo": "Materiais", "camada": "Backend"},
    {"id": 21, "nome": "POST /api/users - Incluir Usuário", "ali": 1, "der": 5, "complexidade": "Baixa", "pf": 3, "modulo": "Usuários", "camada": "Backend"},
    {"id": 22, "nome": "PUT /api/users/{id} - Alterar Usuário", "ali": 1, "der": 5, "complexidade": "Baixa", "pf": 3, "modulo": "Usuários", "camada": "Backend"},
    {"id": 23, "nome": "DELETE /api/users/{id} - Desativar Usuário", "ali": 1, "der": 2, "complexidade": "Baixa", "pf": 3, "modulo": "Usuários", "camada": "Backend"},
    {"id": 24, "nome": "POST /api/users/login - Autenticar Usuário", "ali": 1, "der": 3, "complexidade": "Baixa", "pf": 3, "modulo": "Usuários", "camada": "Backend"},
    {"id": 25, "nome": "PUT /api/settings/parameters - Atualizar Parâmetros", "ali": 1, "der": 8, "complexidade": "Baixa", "pf": 3, "modulo": "Configurações", "camada": "Backend"},
    {"id": 26, "nome": "POST /api/settings/bank-prices - Incluir Banco Preços", "ali": 1, "der": 6, "complexidade": "Baixa", "pf": 3, "modulo": "Configurações", "camada": "Backend"},
    {"id": 27, "nome": "PUT /api/settings/bank-prices/{id} - Alterar Banco Preços", "ali": 1, "der": 6, "complexidade": "Baixa", "pf": 3, "modulo": "Configurações", "camada": "Backend"},
    {"id": 28, "nome": "DELETE /api/settings/bank-prices/{id} - Excluir Banco Preços", "ali": 1, "der": 2, "complexidade": "Baixa", "pf": 3, "modulo": "Configurações", "camada": "Backend"},
    {"id": 29, "nome": "PUT /api/settings/revaluation - Atualizar Fatores Reavaliação", "ali": 1, "der": 5, "complexidade": "Baixa", "pf": 3, "modulo": "Configurações", "camada": "Backend"},
    {"id": 30, "nome": "PUT /api/settings/integrations/{provider} - Configurar Integração", "ali": 1, "der": 4, "complexidade": "Baixa", "pf": 3, "modulo": "Integrações", "camada": "Backend"},
    {"id": 31, "nome": "POST /api/projects/{id}/config - Criar Versão Config", "ali": 2, "der": 22, "complexidade": "Alta", "pf": 6, "modulo": "Configurações", "camada": "Backend"},
    {"id": 32, "nome": "POST /api/financial/api-config - Incluir Config Custo", "ali": 1, "der": 11, "complexidade": "Média", "pf": 4, "modulo": "Financeiro", "camada": "Backend"},
    {"id": 33, "nome": "PATCH /api/financial/api-config/{id} - Alterar Config Custo", "ali": 1, "der": 11, "complexidade": "Média", "pf": 4, "modulo": "Financeiro", "camada": "Backend"},
    {"id": 34, "nome": "POST /api/financial/transactions - Registrar Transação", "ali": 2, "der": 12, "complexidade": "Média", "pf": 4, "modulo": "Financeiro", "camada": "Backend"},
]

# Saídas Externas (SE) - Backend
SE_BACKEND = [
    {"id": 1, "nome": "Celery Task: Processar Cotação (IA + Scraping)", "ali": 4, "der": 25, "complexidade": "Alta", "pf": 7, "modulo": "Cotações", "camada": "Backend"},
    {"id": 2, "nome": "POST /api/quotes/{id}/generate-pdf - Gerar PDF", "ali": 3, "der": 20, "complexidade": "Alta", "pf": 7, "modulo": "Cotações", "camada": "Backend"},
    {"id": 3, "nome": "Calcular Estatísticas Cotação (média, min, max)", "ali": 2, "der": 8, "complexidade": "Média", "pf": 5, "modulo": "Cotações", "camada": "Backend"},
    {"id": 4, "nome": "Detectar Outliers de Preços", "ali": 2, "der": 6, "complexidade": "Média", "pf": 5, "modulo": "Cotações", "camada": "Backend"},
    {"id": 5, "nome": "POST /api/materials/suggest - Sugerir Materiais", "ali": 2, "der": 15, "complexidade": "Alta", "pf": 7, "modulo": "Materiais", "camada": "Backend"},
    {"id": 6, "nome": "GET /api/financial/report - Relatório Financeiro", "ali": 3, "der": 12, "complexidade": "Alta", "pf": 7, "modulo": "Financeiro", "camada": "Backend"},
    {"id": 7, "nome": "GET /api/financial/summary - Dashboard Financeiro", "ali": 2, "der": 8, "complexidade": "Média", "pf": 5, "modulo": "Financeiro", "camada": "Backend"},
    {"id": 8, "nome": "POST /api/settings/integrations/{provider}/test - Testar API", "ali": 2, "der": 4, "complexidade": "Baixa", "pf": 4, "modulo": "Integrações", "camada": "Backend"},
]

# Consultas Externas (CE) - Backend
CE_BACKEND = [
    {"id": 1, "nome": "GET /api/clients/{id} - Consultar Cliente", "ali": 1, "der": 14, "complexidade": "Média", "pf": 4, "modulo": "Clientes", "camada": "Backend"},
    {"id": 2, "nome": "GET /api/clients - Listar Clientes", "ali": 1, "der": 10, "complexidade": "Baixa", "pf": 3, "modulo": "Clientes", "camada": "Backend"},
    {"id": 3, "nome": "GET /api/projects/{id} - Consultar Projeto", "ali": 2, "der": 16, "complexidade": "Média", "pf": 4, "modulo": "Projetos", "camada": "Backend"},
    {"id": 4, "nome": "GET /api/projects - Listar Projetos", "ali": 2, "der": 10, "complexidade": "Média", "pf": 4, "modulo": "Projetos", "camada": "Backend"},
    {"id": 5, "nome": "GET /api/quotes/{id} - Consultar Cotação", "ali": 4, "der": 25, "complexidade": "Alta", "pf": 6, "modulo": "Cotações", "camada": "Backend"},
    {"id": 6, "nome": "GET /api/quotes - Listar Cotações", "ali": 3, "der": 12, "complexidade": "Média", "pf": 4, "modulo": "Cotações", "camada": "Backend"},
    {"id": 7, "nome": "GET /api/quotes/{id}/pdf - Download PDF", "ali": 2, "der": 4, "complexidade": "Baixa", "pf": 3, "modulo": "Cotações", "camada": "Backend"},
    {"id": 8, "nome": "GET /api/quotes/{id}/screenshots/{id} - Download Screenshot", "ali": 1, "der": 4, "complexidade": "Baixa", "pf": 3, "modulo": "Cotações", "camada": "Backend"},
    {"id": 9, "nome": "GET /api/materials/{id} - Consultar Material", "ali": 2, "der": 12, "complexidade": "Média", "pf": 4, "modulo": "Materiais", "camada": "Backend"},
    {"id": 10, "nome": "GET /api/materials - Listar Materiais", "ali": 2, "der": 10, "complexidade": "Média", "pf": 4, "modulo": "Materiais", "camada": "Backend"},
    {"id": 11, "nome": "GET /api/materials/{id}/characteristics - Listar Características", "ali": 2, "der": 6, "complexidade": "Baixa", "pf": 3, "modulo": "Materiais", "camada": "Backend"},
    {"id": 12, "nome": "GET /api/materials/items/{id} - Consultar Item", "ali": 3, "der": 11, "complexidade": "Média", "pf": 4, "modulo": "Itens", "camada": "Backend"},
    {"id": 13, "nome": "GET /api/materials/items/list - Listar Itens", "ali": 3, "der": 10, "complexidade": "Média", "pf": 4, "modulo": "Itens", "camada": "Backend"},
    {"id": 14, "nome": "GET /api/materials/items/status/options - Opções Status", "ali": 1, "der": 3, "complexidade": "Baixa", "pf": 3, "modulo": "Itens", "camada": "Backend"},
    {"id": 15, "nome": "GET /api/materials/options/list - Opções Materiais", "ali": 1, "der": 3, "complexidade": "Baixa", "pf": 3, "modulo": "Materiais", "camada": "Backend"},
    {"id": 16, "nome": "GET /api/materials/options/characteristic-types - Tipos Características", "ali": 1, "der": 3, "complexidade": "Baixa", "pf": 3, "modulo": "Materiais", "camada": "Backend"},
    {"id": 17, "nome": "GET /api/users - Listar Usuários", "ali": 1, "der": 5, "complexidade": "Baixa", "pf": 3, "modulo": "Usuários", "camada": "Backend"},
    {"id": 18, "nome": "GET /api/settings/parameters - Consultar Parâmetros", "ali": 1, "der": 8, "complexidade": "Baixa", "pf": 3, "modulo": "Configurações", "camada": "Backend"},
    {"id": 19, "nome": "GET /api/settings/bank-prices - Listar Banco Preços", "ali": 1, "der": 6, "complexidade": "Baixa", "pf": 3, "modulo": "Configurações", "camada": "Backend"},
    {"id": 20, "nome": "GET /api/settings/revaluation - Consultar Fatores", "ali": 1, "der": 5, "complexidade": "Baixa", "pf": 3, "modulo": "Configurações", "camada": "Backend"},
    {"id": 21, "nome": "GET /api/settings/integrations/{provider} - Consultar Integração", "ali": 1, "der": 4, "complexidade": "Baixa", "pf": 3, "modulo": "Integrações", "camada": "Backend"},
    {"id": 22, "nome": "GET /api/settings/serpapi-locations - Listar Localizações", "ali": 1, "der": 3, "complexidade": "Baixa", "pf": 3, "modulo": "Configurações", "camada": "Backend"},
    {"id": 23, "nome": "GET /api/settings/anthropic-models - Listar Modelos IA", "ali": 1, "der": 3, "complexidade": "Baixa", "pf": 3, "modulo": "Configurações", "camada": "Backend"},
    {"id": 24, "nome": "GET /api/projects/{id}/config - Consultar Config Projeto", "ali": 2, "der": 22, "complexidade": "Alta", "pf": 6, "modulo": "Configurações", "camada": "Backend"},
    {"id": 25, "nome": "GET /api/financial/transactions - Listar Transações", "ali": 2, "der": 12, "complexidade": "Média", "pf": 4, "modulo": "Financeiro", "camada": "Backend"},
    {"id": 26, "nome": "GET /api/financial/api-config - Listar Configs Custo", "ali": 1, "der": 11, "complexidade": "Média", "pf": 4, "modulo": "Financeiro", "camada": "Backend"},
    {"id": 27, "nome": "GET /api/financial/api-config/active - Configs Ativas", "ali": 1, "der": 6, "complexidade": "Baixa", "pf": 3, "modulo": "Financeiro", "camada": "Backend"},
]

# Frontend - Telas/Componentes
FRONTEND_DATA = [
    {"id": 1, "nome": "Página Login (/login)", "tipo": "Tela", "complexidade": "Baixa", "pf": 4, "modulo": "Autenticação", "camada": "Frontend"},
    {"id": 2, "nome": "Página Home (/)", "tipo": "Tela", "complexidade": "Baixa", "pf": 3, "modulo": "Geral", "camada": "Frontend"},
    {"id": 3, "nome": "Página Nova Cotação (/cotacao)", "tipo": "Tela", "complexidade": "Alta", "pf": 8, "modulo": "Cotações", "camada": "Frontend"},
    {"id": 4, "nome": "Página Detalhes Cotação (/cotacao/[id])", "tipo": "Tela", "complexidade": "Alta", "pf": 8, "modulo": "Cotações", "camada": "Frontend"},
    {"id": 5, "nome": "Página Histórico (/historico)", "tipo": "Tela", "complexidade": "Média", "pf": 6, "modulo": "Cotações", "camada": "Frontend"},
    {"id": 6, "nome": "Página Clientes (/cadastros/clientes)", "tipo": "Tela", "complexidade": "Média", "pf": 6, "modulo": "Clientes", "camada": "Frontend"},
    {"id": 7, "nome": "Página Projetos (/cadastros/projetos)", "tipo": "Tela", "complexidade": "Média", "pf": 6, "modulo": "Projetos", "camada": "Frontend"},
    {"id": 8, "nome": "Página Detalhes Projeto (/cadastros/projetos/[id])", "tipo": "Tela", "complexidade": "Alta", "pf": 8, "modulo": "Projetos", "camada": "Frontend"},
    {"id": 9, "nome": "Página Materiais (/cadastros/materiais)", "tipo": "Tela", "complexidade": "Média", "pf": 6, "modulo": "Materiais", "camada": "Frontend"},
    {"id": 10, "nome": "Página Itens (/cadastros/itens)", "tipo": "Tela", "complexidade": "Média", "pf": 6, "modulo": "Itens", "camada": "Frontend"},
    {"id": 11, "nome": "Página Gerar Itens (/cadastros/itens/gerar)", "tipo": "Tela", "complexidade": "Alta", "pf": 8, "modulo": "Itens", "camada": "Frontend"},
    {"id": 12, "nome": "Página Características (/cadastros/caracteristicas)", "tipo": "Tela", "complexidade": "Média", "pf": 5, "modulo": "Materiais", "camada": "Frontend"},
    {"id": 13, "nome": "Página Parâmetros (/configuracoes/parametros)", "tipo": "Tela", "complexidade": "Média", "pf": 5, "modulo": "Configurações", "camada": "Frontend"},
    {"id": 14, "nome": "Página Banco Preços (/configuracoes/banco-precos)", "tipo": "Tela", "complexidade": "Média", "pf": 6, "modulo": "Configurações", "camada": "Frontend"},
    {"id": 15, "nome": "Página Fator Reavaliação (/configuracoes/fator-reavaliacao)", "tipo": "Tela", "complexidade": "Média", "pf": 5, "modulo": "Configurações", "camada": "Frontend"},
    {"id": 16, "nome": "Página Integrações (/configuracoes/integracoes)", "tipo": "Tela", "complexidade": "Média", "pf": 6, "modulo": "Integrações", "camada": "Frontend"},
    {"id": 17, "nome": "Página Usuários (/admin/usuarios)", "tipo": "Tela", "complexidade": "Média", "pf": 6, "modulo": "Usuários", "camada": "Frontend"},
    {"id": 18, "nome": "Página Financeiro (/admin/financeiro)", "tipo": "Tela", "complexidade": "Alta", "pf": 8, "modulo": "Financeiro", "camada": "Frontend"},
    {"id": 19, "nome": "Página Conta (/conta)", "tipo": "Tela", "complexidade": "Baixa", "pf": 4, "modulo": "Usuários", "camada": "Frontend"},
    {"id": 20, "nome": "Componente Sidebar", "tipo": "Componente", "complexidade": "Média", "pf": 4, "modulo": "Geral", "camada": "Frontend"},
    {"id": 21, "nome": "Componente ProtectedRoute", "tipo": "Componente", "complexidade": "Baixa", "pf": 3, "modulo": "Autenticação", "camada": "Frontend"},
    {"id": 22, "nome": "Contexto AuthContext", "tipo": "Componente", "complexidade": "Média", "pf": 4, "modulo": "Autenticação", "camada": "Frontend"},
    {"id": 23, "nome": "Cliente API (lib/api.ts)", "tipo": "Serviço", "complexidade": "Alta", "pf": 6, "modulo": "Geral", "camada": "Frontend"},
]

# Infraestrutura
INFRA_DATA = [
    {"id": 1, "nome": "Dockerfile Backend (FastAPI)", "tipo": "Container", "complexidade": "Média", "pf": 4, "modulo": "Infraestrutura", "camada": "Infra"},
    {"id": 2, "nome": "Dockerfile Frontend (Next.js)", "tipo": "Container", "complexidade": "Média", "pf": 4, "modulo": "Infraestrutura", "camada": "Infra"},
    {"id": 3, "nome": "Docker Compose (Orquestração)", "tipo": "Orquestração", "complexidade": "Alta", "pf": 6, "modulo": "Infraestrutura", "camada": "Infra"},
    {"id": 4, "nome": "Configuração PostgreSQL", "tipo": "Banco de Dados", "complexidade": "Média", "pf": 4, "modulo": "Infraestrutura", "camada": "Infra"},
    {"id": 5, "nome": "Configuração Redis", "tipo": "Cache/Queue", "complexidade": "Baixa", "pf": 3, "modulo": "Infraestrutura", "camada": "Infra"},
    {"id": 6, "nome": "Configuração Celery Worker", "tipo": "Worker", "complexidade": "Média", "pf": 4, "modulo": "Infraestrutura", "camada": "Infra"},
    {"id": 7, "nome": "Alembic Migrations (11 versões)", "tipo": "Migrations", "complexidade": "Alta", "pf": 8, "modulo": "Banco de Dados", "camada": "Infra"},
    {"id": 8, "nome": "Configuração CORS", "tipo": "Segurança", "complexidade": "Baixa", "pf": 2, "modulo": "Segurança", "camada": "Infra"},
    {"id": 9, "nome": "Criptografia de Credenciais", "tipo": "Segurança", "complexidade": "Média", "pf": 4, "modulo": "Segurança", "camada": "Infra"},
    {"id": 10, "nome": "Hash de Senhas (bcrypt)", "tipo": "Segurança", "complexidade": "Baixa", "pf": 2, "modulo": "Segurança", "camada": "Infra"},
    {"id": 11, "nome": "Gestão de Storage (uploads/pdfs)", "tipo": "Storage", "complexidade": "Média", "pf": 4, "modulo": "Infraestrutura", "camada": "Infra"},
    {"id": 12, "nome": "Configuração Playwright (Browser)", "tipo": "Scraping", "complexidade": "Alta", "pf": 5, "modulo": "Infraestrutura", "camada": "Infra"},
]

# Serviços Backend (Lógica de Negócio)
SERVICES_DATA = [
    {"id": 1, "nome": "ClaudeClient - Integração Anthropic", "tipo": "Serviço", "complexidade": "Alta", "pf": 8, "modulo": "Integrações", "camada": "Backend"},
    {"id": 2, "nome": "SearchProvider - Integração SerpAPI", "tipo": "Serviço", "complexidade": "Alta", "pf": 7, "modulo": "Integrações", "camada": "Backend"},
    {"id": 3, "nome": "PriceExtractor - Web Scraping", "tipo": "Serviço", "complexidade": "Alta", "pf": 10, "modulo": "Cotações", "camada": "Backend"},
    {"id": 4, "nome": "PDFGenerator - Geração de Relatórios", "tipo": "Serviço", "complexidade": "Alta", "pf": 8, "modulo": "Cotações", "camada": "Backend"},
    {"id": 5, "nome": "QuoteTasks - Processamento Assíncrono", "tipo": "Serviço", "complexidade": "Alta", "pf": 10, "modulo": "Cotações", "camada": "Backend"},
    {"id": 6, "nome": "Security - Criptografia/Auth", "tipo": "Serviço", "complexidade": "Média", "pf": 5, "modulo": "Segurança", "camada": "Backend"},
]

def set_cell_shading(cell, color):
    """Define a cor de fundo de uma célula"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_word_document():
    """Cria documento Word com visão executiva"""
    doc = Document()

    # Configurar margens
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ========== CAPA ==========
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_heading('ANÁLISE DE PONTOS DE FUNÇÃO', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Sistema de Cotação Automatizada de Preços')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)

    doc.add_paragraph()

    subtitle2 = doc.add_paragraph('Reavaliação Patrimonial para Órgãos Públicos')
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle2.runs[0].font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph()

    # Informações do documento
    info_table = doc.add_table(rows=4, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    info_data = [
        ("Versão:", "1.0"),
        ("Data:", datetime.now().strftime("%d/%m/%Y")),
        ("Metodologia:", "IFPUG CPM 4.3.1"),
        ("Tipo de Contagem:", "Contagem Detalhada"),
    ]

    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value
        info_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True

    doc.add_page_break()

    # ========== SUMÁRIO EXECUTIVO ==========
    doc.add_heading('1. SUMÁRIO EXECUTIVO', level=1)

    doc.add_heading('1.1 Objetivo do Documento', level=2)
    doc.add_paragraph(
        'Este documento apresenta a Análise de Pontos de Função (APF) do Sistema de Cotação '
        'Automatizada de Preços, utilizando a metodologia IFPUG. A contagem considera todas as '
        'funcionalidades do backend, frontend e infraestrutura necessárias para operação do sistema.'
    )

    doc.add_heading('1.2 Visão Geral do Sistema', level=2)
    doc.add_paragraph(
        'O sistema automatiza o processo de pesquisa e cotação de preços de mercado para fins '
        'de reavaliação patrimonial de órgãos públicos brasileiros. Utiliza Inteligência Artificial '
        '(Claude/Anthropic) e Web Scraping para analisar produtos e coletar preços automaticamente.'
    )

    # Tabela de Resultados
    doc.add_heading('1.3 Resultado da Contagem', level=2)

    # Calcular totais
    total_ali = sum(x['pf'] for x in ALI_DATA)
    total_aie = sum(x['pf'] for x in AIE_DATA)
    total_ee = sum(x['pf'] for x in EE_BACKEND)
    total_se = sum(x['pf'] for x in SE_BACKEND)
    total_ce = sum(x['pf'] for x in CE_BACKEND)
    total_frontend = sum(x['pf'] for x in FRONTEND_DATA)
    total_infra = sum(x['pf'] for x in INFRA_DATA)
    total_services = sum(x['pf'] for x in SERVICES_DATA)

    pf_brutos = total_ali + total_aie + total_ee + total_se + total_ce
    pf_frontend_infra = total_frontend + total_infra + total_services
    pf_total = pf_brutos + pf_frontend_infra

    # VAF = 1.11 (calculado anteriormente)
    vaf = 1.11
    pf_ajustados = round(pf_total * vaf)

    result_table = doc.add_table(rows=5, cols=2)
    result_table.style = 'Table Grid'

    result_data = [
        ("Pontos de Função Brutos (Backend)", f"{pf_brutos} PF"),
        ("Pontos de Função (Frontend + Infra + Serviços)", f"{pf_frontend_infra} PF"),
        ("Total Pontos de Função Brutos", f"{pf_total} PF"),
        ("Fator de Ajuste (VAF)", f"{vaf}"),
        ("PONTOS DE FUNÇÃO AJUSTADOS", f"{pf_ajustados} PF"),
    ]

    for i, (label, value) in enumerate(result_data):
        result_table.rows[i].cells[0].text = label
        result_table.rows[i].cells[1].text = value
        if i == 4:  # Última linha em destaque
            result_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
            result_table.rows[i].cells[1].paragraphs[0].runs[0].font.bold = True
            set_cell_shading(result_table.rows[i].cells[0], 'D9E2F3')
            set_cell_shading(result_table.rows[i].cells[1], 'D9E2F3')

    doc.add_paragraph()

    # ========== DISTRIBUIÇÃO POR CAMADA ==========
    doc.add_heading('1.4 Distribuição por Camada', level=2)

    layer_table = doc.add_table(rows=5, cols=3)
    layer_table.style = 'Table Grid'

    # Cabeçalho
    header_cells = layer_table.rows[0].cells
    header_cells[0].text = "Camada"
    header_cells[1].text = "Pontos de Função"
    header_cells[2].text = "Percentual"
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    layer_data = [
        ("Backend (APIs + ALI/AIE)", pf_brutos),
        ("Backend (Serviços)", total_services),
        ("Frontend (Telas)", total_frontend),
        ("Infraestrutura", total_infra),
    ]

    for i, (camada, pf) in enumerate(layer_data, 1):
        layer_table.rows[i].cells[0].text = camada
        layer_table.rows[i].cells[1].text = str(pf)
        layer_table.rows[i].cells[2].text = f"{(pf/pf_total)*100:.1f}%"

    doc.add_paragraph()

    # ========== ESTIMATIVA DE ESFORÇO ==========
    doc.add_heading('1.5 Estimativa de Esforço', level=2)

    doc.add_paragraph(
        'Com base nos Pontos de Função ajustados e produtividade média de mercado:'
    )

    effort_table = doc.add_table(rows=4, cols=4)
    effort_table.style = 'Table Grid'

    # Cabeçalho
    headers = ["Cenário", "Horas/PF", "Esforço Total", "Homem-Mês"]
    for i, header in enumerate(headers):
        effort_table.rows[0].cells[i].text = header
        effort_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        set_cell_shading(effort_table.rows[0].cells[i], '4472C4')
        effort_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    effort_data = [
        ("Otimista", 10, pf_ajustados * 10, pf_ajustados * 10 / 176),
        ("Realista", 14, pf_ajustados * 14, pf_ajustados * 14 / 176),
        ("Pessimista", 18, pf_ajustados * 18, pf_ajustados * 18 / 176),
    ]

    for i, (cenario, hpf, esforco, hm) in enumerate(effort_data, 1):
        effort_table.rows[i].cells[0].text = cenario
        effort_table.rows[i].cells[1].text = str(hpf)
        effort_table.rows[i].cells[2].text = f"{esforco:,.0f} h"
        effort_table.rows[i].cells[3].text = f"{hm:.1f} HM"

    doc.add_paragraph()

    # ========== ESTIMATIVA DE CUSTO ==========
    doc.add_heading('1.6 Estimativa de Custo', level=2)

    doc.add_paragraph(
        'Considerando valores médios de mercado para equipe de desenvolvimento:'
    )

    cost_table = doc.add_table(rows=5, cols=3)
    cost_table.style = 'Table Grid'

    # Cabeçalho
    cost_headers = ["Componente", "Cenário Realista", "Observação"]
    for i, header in enumerate(cost_headers):
        cost_table.rows[0].cells[i].text = header
        cost_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cost_table.rows[0].cells[i], '4472C4')
        cost_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    horas_realista = pf_ajustados * 14
    custo_hora = 150  # R$ 150/hora média
    custo_total = horas_realista * custo_hora

    cost_data = [
        ("Esforço", f"{horas_realista:,.0f} horas", f"{horas_realista/176:.1f} Homem-Mês"),
        ("Custo Hora Médio", f"R$ {custo_hora},00", "Média mercado BR"),
        ("Custo Desenvolvimento", f"R$ {custo_total:,.2f}", "Equipe multidisciplinar"),
        ("Custo Total Estimado", f"R$ {custo_total*1.2:,.2f}", "+20% contingência"),
    ]

    for i, (comp, valor, obs) in enumerate(cost_data, 1):
        cost_table.rows[i].cells[0].text = comp
        cost_table.rows[i].cells[1].text = valor
        cost_table.rows[i].cells[2].text = obs
        if i == 4:
            for cell in cost_table.rows[i].cells:
                cell.paragraphs[0].runs[0].font.bold = True
                set_cell_shading(cell, 'E2EFDA')

    doc.add_page_break()

    # ========== DETALHAMENTO POR TIPO ==========
    doc.add_heading('2. DETALHAMENTO DA CONTAGEM', level=1)

    doc.add_heading('2.1 Arquivos Lógicos Internos (ALI)', level=2)
    doc.add_paragraph(f'Total: {len(ALI_DATA)} entidades = {total_ali} PF')

    ali_table = doc.add_table(rows=len(ALI_DATA)+1, cols=4)
    ali_table.style = 'Table Grid'

    ali_headers = ["Entidade", "DER", "Complexidade", "PF"]
    for i, header in enumerate(ali_headers):
        ali_table.rows[0].cells[i].text = header
        ali_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        set_cell_shading(ali_table.rows[0].cells[i], '4472C4')
        ali_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for i, item in enumerate(ALI_DATA, 1):
        ali_table.rows[i].cells[0].text = item['nome']
        ali_table.rows[i].cells[1].text = str(item['der'])
        ali_table.rows[i].cells[2].text = item['complexidade']
        ali_table.rows[i].cells[3].text = str(item['pf'])

    doc.add_paragraph()

    doc.add_heading('2.2 Arquivos de Interface Externa (AIE)', level=2)
    doc.add_paragraph(f'Total: {len(AIE_DATA)} integrações = {total_aie} PF')

    aie_table = doc.add_table(rows=len(AIE_DATA)+1, cols=4)
    aie_table.style = 'Table Grid'

    for i, header in enumerate(ali_headers):
        aie_table.rows[0].cells[i].text = header
        aie_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        set_cell_shading(aie_table.rows[0].cells[i], '4472C4')
        aie_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for i, item in enumerate(AIE_DATA, 1):
        aie_table.rows[i].cells[0].text = item['nome']
        aie_table.rows[i].cells[1].text = str(item['der'])
        aie_table.rows[i].cells[2].text = item['complexidade']
        aie_table.rows[i].cells[3].text = str(item['pf'])

    doc.add_paragraph()

    # Resumo de Transações
    doc.add_heading('2.3 Funções de Transação (Backend)', level=2)

    trans_summary = doc.add_table(rows=4, cols=3)
    trans_summary.style = 'Table Grid'

    trans_headers = ["Tipo", "Quantidade", "Pontos de Função"]
    for i, header in enumerate(trans_headers):
        trans_summary.rows[0].cells[i].text = header
        trans_summary.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        set_cell_shading(trans_summary.rows[0].cells[i], '4472C4')
        trans_summary.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    trans_data = [
        ("Entradas Externas (EE)", len(EE_BACKEND), total_ee),
        ("Saídas Externas (SE)", len(SE_BACKEND), total_se),
        ("Consultas Externas (CE)", len(CE_BACKEND), total_ce),
    ]

    for i, (tipo, qtd, pf) in enumerate(trans_data, 1):
        trans_summary.rows[i].cells[0].text = tipo
        trans_summary.rows[i].cells[1].text = str(qtd)
        trans_summary.rows[i].cells[2].text = str(pf)

    doc.add_paragraph()

    doc.add_heading('2.4 Frontend e Infraestrutura', level=2)

    front_summary = doc.add_table(rows=4, cols=3)
    front_summary.style = 'Table Grid'

    for i, header in enumerate(trans_headers):
        front_summary.rows[0].cells[i].text = header
        front_summary.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        set_cell_shading(front_summary.rows[0].cells[i], '4472C4')
        front_summary.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    front_data = [
        ("Telas e Componentes Frontend", len(FRONTEND_DATA), total_frontend),
        ("Serviços Backend", len(SERVICES_DATA), total_services),
        ("Componentes Infraestrutura", len(INFRA_DATA), total_infra),
    ]

    for i, (tipo, qtd, pf) in enumerate(front_data, 1):
        front_summary.rows[i].cells[0].text = tipo
        front_summary.rows[i].cells[1].text = str(qtd)
        front_summary.rows[i].cells[2].text = str(pf)

    doc.add_page_break()

    # ========== STACK TECNOLÓGICO ==========
    doc.add_heading('3. STACK TECNOLÓGICO', level=1)

    doc.add_heading('3.1 Backend', level=2)
    backend_stack = [
        "Python 3.11+ (Linguagem principal)",
        "FastAPI 0.109.0 (Framework web)",
        "SQLAlchemy 2.0.25 (ORM)",
        "PostgreSQL 15 (Banco de dados)",
        "Celery 5.3.6 (Processamento assíncrono)",
        "Redis 7 (Message broker)",
        "Playwright 1.41.0 (Web scraping)",
        "Anthropic SDK 0.18.1 (Claude AI)",
    ]
    for item in backend_stack:
        doc.add_paragraph(f"• {item}", style='List Bullet')

    doc.add_heading('3.2 Frontend', level=2)
    frontend_stack = [
        "Next.js 14.1.0 (Framework React)",
        "React 18.2.0 (UI Library)",
        "TypeScript 5.3.3 (Linguagem)",
        "Tailwind CSS 3.4.1 (Estilização)",
        "Axios 1.6.7 (Cliente HTTP)",
        "SWR 2.2.5 (Data fetching)",
    ]
    for item in frontend_stack:
        doc.add_paragraph(f"• {item}", style='List Bullet')

    doc.add_heading('3.3 Infraestrutura', level=2)
    infra_stack = [
        "Docker + Docker Compose (Containerização)",
        "Alembic (Migrations de banco)",
        "bcrypt (Hash de senhas)",
        "cryptography (Criptografia de credenciais)",
    ]
    for item in infra_stack:
        doc.add_paragraph(f"• {item}", style='List Bullet')

    doc.add_page_break()

    # ========== GLOSSÁRIO ==========
    doc.add_heading('4. GLOSSÁRIO', level=1)

    glossary = [
        ("ALI", "Arquivo Lógico Interno - Dados mantidos pela aplicação"),
        ("AIE", "Arquivo de Interface Externa - Dados de sistemas externos"),
        ("CE", "Consulta Externa - Recuperação de dados sem processamento"),
        ("DER", "Dados Elementares Referenciados - Campos da função"),
        ("EE", "Entrada Externa - Processo que mantém dados"),
        ("IFPUG", "International Function Point Users Group"),
        ("PF", "Ponto de Função - Unidade de medida"),
        ("RLR", "Registro Lógico Referenciado - Subgrupos de dados"),
        ("SE", "Saída Externa - Processo com cálculos/derivações"),
        ("VAF", "Value Adjustment Factor - Fator de ajuste"),
    ]

    gloss_table = doc.add_table(rows=len(glossary)+1, cols=2)
    gloss_table.style = 'Table Grid'

    gloss_table.rows[0].cells[0].text = "Sigla"
    gloss_table.rows[0].cells[1].text = "Definição"
    for cell in gloss_table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for i, (sigla, definicao) in enumerate(glossary, 1):
        gloss_table.rows[i].cells[0].text = sigla
        gloss_table.rows[i].cells[1].text = definicao

    # Salvar documento
    doc.save('C:/Projeto_reavaliacao/APF_Sistema_Cotacao.docx')
    print("Documento Word gerado: APF_Sistema_Cotacao.docx")


def create_excel_spreadsheet():
    """Cria planilha Excel com detalhamento completo"""
    wb = openpyxl.Workbook()

    # Estilos
    header_font = Font(bold=True, color='FFFFFF', size=11)
    header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
    alt_fill = PatternFill(start_color='D9E2F3', end_color='D9E2F3', fill_type='solid')
    total_fill = PatternFill(start_color='E2EFDA', end_color='E2EFDA', fill_type='solid')
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    center_align = Alignment(horizontal='center', vertical='center')
    wrap_align = Alignment(horizontal='left', vertical='center', wrap_text=True)

    # ========== ABA RESUMO ==========
    ws_resumo = wb.active
    ws_resumo.title = "Resumo Executivo"

    # Título
    ws_resumo['A1'] = "ANÁLISE DE PONTOS DE FUNÇÃO - RESUMO EXECUTIVO"
    ws_resumo['A1'].font = Font(bold=True, size=16)
    ws_resumo.merge_cells('A1:F1')

    ws_resumo['A2'] = "Sistema de Cotação Automatizada de Preços"
    ws_resumo['A2'].font = Font(size=12)
    ws_resumo.merge_cells('A2:F2')

    ws_resumo['A3'] = f"Data: {datetime.now().strftime('%d/%m/%Y')} | Metodologia: IFPUG CPM 4.3.1"
    ws_resumo.merge_cells('A3:F3')

    # Totais
    total_ali = sum(x['pf'] for x in ALI_DATA)
    total_aie = sum(x['pf'] for x in AIE_DATA)
    total_ee = sum(x['pf'] for x in EE_BACKEND)
    total_se = sum(x['pf'] for x in SE_BACKEND)
    total_ce = sum(x['pf'] for x in CE_BACKEND)
    total_frontend = sum(x['pf'] for x in FRONTEND_DATA)
    total_infra = sum(x['pf'] for x in INFRA_DATA)
    total_services = sum(x['pf'] for x in SERVICES_DATA)

    pf_brutos_backend = total_ali + total_aie + total_ee + total_se + total_ce
    pf_total = pf_brutos_backend + total_frontend + total_infra + total_services
    vaf = 1.11
    pf_ajustados = round(pf_total * vaf)

    # Tabela de resumo
    row = 5
    headers = ["Categoria", "Tipo", "Quantidade", "Pontos de Função", "% do Total"]
    for col, header in enumerate(headers, 1):
        cell = ws_resumo.cell(row=row, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = thin_border
        cell.alignment = center_align

    data = [
        ("Backend", "ALI - Arquivos Lógicos Internos", len(ALI_DATA), total_ali),
        ("Backend", "AIE - Arquivos Interface Externa", len(AIE_DATA), total_aie),
        ("Backend", "EE - Entradas Externas", len(EE_BACKEND), total_ee),
        ("Backend", "SE - Saídas Externas", len(SE_BACKEND), total_se),
        ("Backend", "CE - Consultas Externas", len(CE_BACKEND), total_ce),
        ("Backend", "Serviços de Negócio", len(SERVICES_DATA), total_services),
        ("Frontend", "Telas e Componentes", len(FRONTEND_DATA), total_frontend),
        ("Infraestrutura", "Componentes de Infra", len(INFRA_DATA), total_infra),
    ]

    for i, (cat, tipo, qtd, pf) in enumerate(data):
        row += 1
        ws_resumo.cell(row=row, column=1, value=cat).border = thin_border
        ws_resumo.cell(row=row, column=2, value=tipo).border = thin_border
        ws_resumo.cell(row=row, column=3, value=qtd).border = thin_border
        ws_resumo.cell(row=row, column=3).alignment = center_align
        ws_resumo.cell(row=row, column=4, value=pf).border = thin_border
        ws_resumo.cell(row=row, column=4).alignment = center_align
        ws_resumo.cell(row=row, column=5, value=f"{(pf/pf_total)*100:.1f}%").border = thin_border
        ws_resumo.cell(row=row, column=5).alignment = center_align
        if i % 2 == 1:
            for col in range(1, 6):
                ws_resumo.cell(row=row, column=col).fill = alt_fill

    # Total
    row += 1
    ws_resumo.cell(row=row, column=1, value="TOTAL BRUTO").font = Font(bold=True)
    ws_resumo.cell(row=row, column=4, value=pf_total).font = Font(bold=True)
    for col in range(1, 6):
        ws_resumo.cell(row=row, column=col).fill = total_fill
        ws_resumo.cell(row=row, column=col).border = thin_border

    row += 2
    ws_resumo.cell(row=row, column=1, value="Fator de Ajuste (VAF)")
    ws_resumo.cell(row=row, column=4, value=vaf)

    row += 1
    ws_resumo.cell(row=row, column=1, value="PONTOS DE FUNÇÃO AJUSTADOS").font = Font(bold=True, size=12)
    ws_resumo.cell(row=row, column=4, value=pf_ajustados).font = Font(bold=True, size=12)
    for col in range(1, 6):
        ws_resumo.cell(row=row, column=col).fill = PatternFill(start_color='92D050', end_color='92D050', fill_type='solid')

    # Estimativa de Esforço
    row += 3
    ws_resumo.cell(row=row, column=1, value="ESTIMATIVA DE ESFORÇO E CUSTO").font = Font(bold=True, size=12)
    ws_resumo.merge_cells(f'A{row}:E{row}')

    row += 1
    for col, header in enumerate(["Cenário", "Horas/PF", "Esforço (h)", "Homem-Mês", "Custo Estimado (R$)"], 1):
        cell = ws_resumo.cell(row=row, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = thin_border
        cell.alignment = center_align

    custo_hora = 150
    scenarios = [
        ("Otimista", 10),
        ("Realista", 14),
        ("Pessimista", 18),
    ]

    for cenario, hpf in scenarios:
        row += 1
        esforco = pf_ajustados * hpf
        hm = esforco / 176
        custo = esforco * custo_hora

        ws_resumo.cell(row=row, column=1, value=cenario).border = thin_border
        ws_resumo.cell(row=row, column=2, value=hpf).border = thin_border
        ws_resumo.cell(row=row, column=2).alignment = center_align
        ws_resumo.cell(row=row, column=3, value=esforco).border = thin_border
        ws_resumo.cell(row=row, column=3).number_format = '#,##0'
        ws_resumo.cell(row=row, column=3).alignment = center_align
        ws_resumo.cell(row=row, column=4, value=round(hm, 1)).border = thin_border
        ws_resumo.cell(row=row, column=4).alignment = center_align
        ws_resumo.cell(row=row, column=5, value=custo).border = thin_border
        ws_resumo.cell(row=row, column=5).number_format = 'R$ #,##0.00'

        if cenario == "Realista":
            for col in range(1, 6):
                ws_resumo.cell(row=row, column=col).fill = alt_fill

    # Ajustar larguras
    ws_resumo.column_dimensions['A'].width = 18
    ws_resumo.column_dimensions['B'].width = 35
    ws_resumo.column_dimensions['C'].width = 15
    ws_resumo.column_dimensions['D'].width = 18
    ws_resumo.column_dimensions['E'].width = 22

    # ========== ABA ALI ==========
    ws_ali = wb.create_sheet("ALI - Dados Internos")

    headers = ["#", "Entidade", "Descrição/Módulo", "DER", "RLR", "Complexidade", "PF"]
    for col, header in enumerate(headers, 1):
        cell = ws_ali.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = thin_border
        cell.alignment = center_align

    for i, item in enumerate(ALI_DATA, 1):
        row = i + 1
        ws_ali.cell(row=row, column=1, value=item['id']).border = thin_border
        ws_ali.cell(row=row, column=1).alignment = center_align
        ws_ali.cell(row=row, column=2, value=item['nome']).border = thin_border
        ws_ali.cell(row=row, column=3, value=item['modulo']).border = thin_border
        ws_ali.cell(row=row, column=4, value=item['der']).border = thin_border
        ws_ali.cell(row=row, column=4).alignment = center_align
        ws_ali.cell(row=row, column=5, value=item['rlr']).border = thin_border
        ws_ali.cell(row=row, column=5).alignment = center_align
        ws_ali.cell(row=row, column=6, value=item['complexidade']).border = thin_border
        ws_ali.cell(row=row, column=6).alignment = center_align
        ws_ali.cell(row=row, column=7, value=item['pf']).border = thin_border
        ws_ali.cell(row=row, column=7).alignment = center_align
        if i % 2 == 0:
            for col in range(1, 8):
                ws_ali.cell(row=row, column=col).fill = alt_fill

    # Total
    row = len(ALI_DATA) + 2
    ws_ali.cell(row=row, column=6, value="TOTAL").font = Font(bold=True)
    ws_ali.cell(row=row, column=7, value=total_ali).font = Font(bold=True)
    for col in range(1, 8):
        ws_ali.cell(row=row, column=col).fill = total_fill
        ws_ali.cell(row=row, column=col).border = thin_border

    ws_ali.column_dimensions['A'].width = 5
    ws_ali.column_dimensions['B'].width = 35
    ws_ali.column_dimensions['C'].width = 18
    ws_ali.column_dimensions['D'].width = 8
    ws_ali.column_dimensions['E'].width = 8
    ws_ali.column_dimensions['F'].width = 14
    ws_ali.column_dimensions['G'].width = 8

    # ========== ABA AIE ==========
    ws_aie = wb.create_sheet("AIE - Dados Externos")

    for col, header in enumerate(headers, 1):
        cell = ws_aie.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = thin_border
        cell.alignment = center_align

    for i, item in enumerate(AIE_DATA, 1):
        row = i + 1
        ws_aie.cell(row=row, column=1, value=item['id']).border = thin_border
        ws_aie.cell(row=row, column=2, value=item['nome']).border = thin_border
        ws_aie.cell(row=row, column=3, value=item['modulo']).border = thin_border
        ws_aie.cell(row=row, column=4, value=item['der']).border = thin_border
        ws_aie.cell(row=row, column=5, value=item['rlr']).border = thin_border
        ws_aie.cell(row=row, column=6, value=item['complexidade']).border = thin_border
        ws_aie.cell(row=row, column=7, value=item['pf']).border = thin_border
        for col in range(1, 8):
            ws_aie.cell(row=row, column=col).alignment = center_align if col != 2 else wrap_align

    row = len(AIE_DATA) + 2
    ws_aie.cell(row=row, column=6, value="TOTAL").font = Font(bold=True)
    ws_aie.cell(row=row, column=7, value=total_aie).font = Font(bold=True)
    for col in range(1, 8):
        ws_aie.cell(row=row, column=col).fill = total_fill
        ws_aie.cell(row=row, column=col).border = thin_border

    for col in ['A', 'B', 'C', 'D', 'E', 'F', 'G']:
        ws_aie.column_dimensions[col].width = ws_ali.column_dimensions[col].width

    # ========== ABA EE (Entradas Externas) ==========
    ws_ee = wb.create_sheet("EE - Entradas Backend")

    headers = ["#", "Endpoint / Função", "Módulo", "ALI Ref", "DER", "Complexidade", "PF"]
    for col, header in enumerate(headers, 1):
        cell = ws_ee.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = thin_border
        cell.alignment = center_align

    for i, item in enumerate(EE_BACKEND, 1):
        row = i + 1
        ws_ee.cell(row=row, column=1, value=item['id']).border = thin_border
        ws_ee.cell(row=row, column=1).alignment = center_align
        ws_ee.cell(row=row, column=2, value=item['nome']).border = thin_border
        ws_ee.cell(row=row, column=3, value=item['modulo']).border = thin_border
        ws_ee.cell(row=row, column=4, value=item['ali']).border = thin_border
        ws_ee.cell(row=row, column=4).alignment = center_align
        ws_ee.cell(row=row, column=5, value=item['der']).border = thin_border
        ws_ee.cell(row=row, column=5).alignment = center_align
        ws_ee.cell(row=row, column=6, value=item['complexidade']).border = thin_border
        ws_ee.cell(row=row, column=6).alignment = center_align
        ws_ee.cell(row=row, column=7, value=item['pf']).border = thin_border
        ws_ee.cell(row=row, column=7).alignment = center_align
        if i % 2 == 0:
            for col in range(1, 8):
                ws_ee.cell(row=row, column=col).fill = alt_fill

    row = len(EE_BACKEND) + 2
    ws_ee.cell(row=row, column=6, value="TOTAL").font = Font(bold=True)
    ws_ee.cell(row=row, column=7, value=total_ee).font = Font(bold=True)
    for col in range(1, 8):
        ws_ee.cell(row=row, column=col).fill = total_fill
        ws_ee.cell(row=row, column=col).border = thin_border

    ws_ee.column_dimensions['A'].width = 5
    ws_ee.column_dimensions['B'].width = 55
    ws_ee.column_dimensions['C'].width = 15
    ws_ee.column_dimensions['D'].width = 10
    ws_ee.column_dimensions['E'].width = 8
    ws_ee.column_dimensions['F'].width = 14
    ws_ee.column_dimensions['G'].width = 8

    # ========== ABA SE (Saídas Externas) ==========
    ws_se = wb.create_sheet("SE - Saídas Backend")

    for col, header in enumerate(headers, 1):
        cell = ws_se.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = thin_border
        cell.alignment = center_align

    for i, item in enumerate(SE_BACKEND, 1):
        row = i + 1
        ws_se.cell(row=row, column=1, value=item['id']).border = thin_border
        ws_se.cell(row=row, column=1).alignment = center_align
        ws_se.cell(row=row, column=2, value=item['nome']).border = thin_border
        ws_se.cell(row=row, column=3, value=item['modulo']).border = thin_border
        ws_se.cell(row=row, column=4, value=item['ali']).border = thin_border
        ws_se.cell(row=row, column=4).alignment = center_align
        ws_se.cell(row=row, column=5, value=item['der']).border = thin_border
        ws_se.cell(row=row, column=5).alignment = center_align
        ws_se.cell(row=row, column=6, value=item['complexidade']).border = thin_border
        ws_se.cell(row=row, column=6).alignment = center_align
        ws_se.cell(row=row, column=7, value=item['pf']).border = thin_border
        ws_se.cell(row=row, column=7).alignment = center_align
        if i % 2 == 0:
            for col in range(1, 8):
                ws_se.cell(row=row, column=col).fill = alt_fill

    row = len(SE_BACKEND) + 2
    ws_se.cell(row=row, column=6, value="TOTAL").font = Font(bold=True)
    ws_se.cell(row=row, column=7, value=total_se).font = Font(bold=True)
    for col in range(1, 8):
        ws_se.cell(row=row, column=col).fill = total_fill
        ws_se.cell(row=row, column=col).border = thin_border

    for col in ['A', 'B', 'C', 'D', 'E', 'F', 'G']:
        ws_se.column_dimensions[col].width = ws_ee.column_dimensions[col].width

    # ========== ABA CE (Consultas Externas) ==========
    ws_ce = wb.create_sheet("CE - Consultas Backend")

    for col, header in enumerate(headers, 1):
        cell = ws_ce.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = thin_border
        cell.alignment = center_align

    for i, item in enumerate(CE_BACKEND, 1):
        row = i + 1
        ws_ce.cell(row=row, column=1, value=item['id']).border = thin_border
        ws_ce.cell(row=row, column=1).alignment = center_align
        ws_ce.cell(row=row, column=2, value=item['nome']).border = thin_border
        ws_ce.cell(row=row, column=3, value=item['modulo']).border = thin_border
        ws_ce.cell(row=row, column=4, value=item['ali']).border = thin_border
        ws_ce.cell(row=row, column=4).alignment = center_align
        ws_ce.cell(row=row, column=5, value=item['der']).border = thin_border
        ws_ce.cell(row=row, column=5).alignment = center_align
        ws_ce.cell(row=row, column=6, value=item['complexidade']).border = thin_border
        ws_ce.cell(row=row, column=6).alignment = center_align
        ws_ce.cell(row=row, column=7, value=item['pf']).border = thin_border
        ws_ce.cell(row=row, column=7).alignment = center_align
        if i % 2 == 0:
            for col in range(1, 8):
                ws_ce.cell(row=row, column=col).fill = alt_fill

    row = len(CE_BACKEND) + 2
    ws_ce.cell(row=row, column=6, value="TOTAL").font = Font(bold=True)
    ws_ce.cell(row=row, column=7, value=total_ce).font = Font(bold=True)
    for col in range(1, 8):
        ws_ce.cell(row=row, column=col).fill = total_fill
        ws_ce.cell(row=row, column=col).border = thin_border

    for col in ['A', 'B', 'C', 'D', 'E', 'F', 'G']:
        ws_ce.column_dimensions[col].width = ws_ee.column_dimensions[col].width

    # ========== ABA FRONTEND ==========
    ws_front = wb.create_sheet("Frontend - Telas")

    headers = ["#", "Componente / Tela", "Tipo", "Módulo", "Complexidade", "PF"]
    for col, header in enumerate(headers, 1):
        cell = ws_front.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = thin_border
        cell.alignment = center_align

    for i, item in enumerate(FRONTEND_DATA, 1):
        row = i + 1
        ws_front.cell(row=row, column=1, value=item['id']).border = thin_border
        ws_front.cell(row=row, column=1).alignment = center_align
        ws_front.cell(row=row, column=2, value=item['nome']).border = thin_border
        ws_front.cell(row=row, column=3, value=item['tipo']).border = thin_border
        ws_front.cell(row=row, column=3).alignment = center_align
        ws_front.cell(row=row, column=4, value=item['modulo']).border = thin_border
        ws_front.cell(row=row, column=5, value=item['complexidade']).border = thin_border
        ws_front.cell(row=row, column=5).alignment = center_align
        ws_front.cell(row=row, column=6, value=item['pf']).border = thin_border
        ws_front.cell(row=row, column=6).alignment = center_align
        if i % 2 == 0:
            for col in range(1, 7):
                ws_front.cell(row=row, column=col).fill = alt_fill

    row = len(FRONTEND_DATA) + 2
    ws_front.cell(row=row, column=5, value="TOTAL").font = Font(bold=True)
    ws_front.cell(row=row, column=6, value=total_frontend).font = Font(bold=True)
    for col in range(1, 7):
        ws_front.cell(row=row, column=col).fill = total_fill
        ws_front.cell(row=row, column=col).border = thin_border

    ws_front.column_dimensions['A'].width = 5
    ws_front.column_dimensions['B'].width = 45
    ws_front.column_dimensions['C'].width = 15
    ws_front.column_dimensions['D'].width = 15
    ws_front.column_dimensions['E'].width = 14
    ws_front.column_dimensions['F'].width = 8

    # ========== ABA SERVIÇOS ==========
    ws_serv = wb.create_sheet("Serviços Backend")

    for col, header in enumerate(headers, 1):
        cell = ws_serv.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = thin_border
        cell.alignment = center_align

    for i, item in enumerate(SERVICES_DATA, 1):
        row = i + 1
        ws_serv.cell(row=row, column=1, value=item['id']).border = thin_border
        ws_serv.cell(row=row, column=1).alignment = center_align
        ws_serv.cell(row=row, column=2, value=item['nome']).border = thin_border
        ws_serv.cell(row=row, column=3, value=item['tipo']).border = thin_border
        ws_serv.cell(row=row, column=3).alignment = center_align
        ws_serv.cell(row=row, column=4, value=item['modulo']).border = thin_border
        ws_serv.cell(row=row, column=5, value=item['complexidade']).border = thin_border
        ws_serv.cell(row=row, column=5).alignment = center_align
        ws_serv.cell(row=row, column=6, value=item['pf']).border = thin_border
        ws_serv.cell(row=row, column=6).alignment = center_align
        if i % 2 == 0:
            for col in range(1, 7):
                ws_serv.cell(row=row, column=col).fill = alt_fill

    row = len(SERVICES_DATA) + 2
    ws_serv.cell(row=row, column=5, value="TOTAL").font = Font(bold=True)
    ws_serv.cell(row=row, column=6, value=total_services).font = Font(bold=True)
    for col in range(1, 7):
        ws_serv.cell(row=row, column=col).fill = total_fill
        ws_serv.cell(row=row, column=col).border = thin_border

    for col in ['A', 'B', 'C', 'D', 'E', 'F']:
        ws_serv.column_dimensions[col].width = ws_front.column_dimensions[col].width

    # ========== ABA INFRAESTRUTURA ==========
    ws_infra = wb.create_sheet("Infraestrutura")

    for col, header in enumerate(headers, 1):
        cell = ws_infra.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = thin_border
        cell.alignment = center_align

    for i, item in enumerate(INFRA_DATA, 1):
        row = i + 1
        ws_infra.cell(row=row, column=1, value=item['id']).border = thin_border
        ws_infra.cell(row=row, column=1).alignment = center_align
        ws_infra.cell(row=row, column=2, value=item['nome']).border = thin_border
        ws_infra.cell(row=row, column=3, value=item['tipo']).border = thin_border
        ws_infra.cell(row=row, column=3).alignment = center_align
        ws_infra.cell(row=row, column=4, value=item['modulo']).border = thin_border
        ws_infra.cell(row=row, column=5, value=item['complexidade']).border = thin_border
        ws_infra.cell(row=row, column=5).alignment = center_align
        ws_infra.cell(row=row, column=6, value=item['pf']).border = thin_border
        ws_infra.cell(row=row, column=6).alignment = center_align
        if i % 2 == 0:
            for col in range(1, 7):
                ws_infra.cell(row=row, column=col).fill = alt_fill

    row = len(INFRA_DATA) + 2
    ws_infra.cell(row=row, column=5, value="TOTAL").font = Font(bold=True)
    ws_infra.cell(row=row, column=6, value=total_infra).font = Font(bold=True)
    for col in range(1, 7):
        ws_infra.cell(row=row, column=col).fill = total_fill
        ws_infra.cell(row=row, column=col).border = thin_border

    for col in ['A', 'B', 'C', 'D', 'E', 'F']:
        ws_infra.column_dimensions[col].width = ws_front.column_dimensions[col].width

    # ========== ABA CONSOLIDADO ==========
    ws_cons = wb.create_sheet("Consolidado Completo")

    headers = ["#", "Camada", "Tipo", "Atividade", "Módulo", "Complexidade", "PF"]
    for col, header in enumerate(headers, 1):
        cell = ws_cons.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = thin_border
        cell.alignment = center_align

    row = 2
    counter = 1

    # ALI
    for item in ALI_DATA:
        ws_cons.cell(row=row, column=1, value=counter).border = thin_border
        ws_cons.cell(row=row, column=2, value="Backend").border = thin_border
        ws_cons.cell(row=row, column=3, value="ALI").border = thin_border
        ws_cons.cell(row=row, column=4, value=item['nome']).border = thin_border
        ws_cons.cell(row=row, column=5, value=item['modulo']).border = thin_border
        ws_cons.cell(row=row, column=6, value=item['complexidade']).border = thin_border
        ws_cons.cell(row=row, column=7, value=item['pf']).border = thin_border
        for col in [1, 2, 3, 6, 7]:
            ws_cons.cell(row=row, column=col).alignment = center_align
        row += 1
        counter += 1

    # AIE
    for item in AIE_DATA:
        ws_cons.cell(row=row, column=1, value=counter).border = thin_border
        ws_cons.cell(row=row, column=2, value="Backend").border = thin_border
        ws_cons.cell(row=row, column=3, value="AIE").border = thin_border
        ws_cons.cell(row=row, column=4, value=item['nome']).border = thin_border
        ws_cons.cell(row=row, column=5, value=item['modulo']).border = thin_border
        ws_cons.cell(row=row, column=6, value=item['complexidade']).border = thin_border
        ws_cons.cell(row=row, column=7, value=item['pf']).border = thin_border
        for col in [1, 2, 3, 6, 7]:
            ws_cons.cell(row=row, column=col).alignment = center_align
        row += 1
        counter += 1

    # EE
    for item in EE_BACKEND:
        ws_cons.cell(row=row, column=1, value=counter).border = thin_border
        ws_cons.cell(row=row, column=2, value="Backend").border = thin_border
        ws_cons.cell(row=row, column=3, value="EE").border = thin_border
        ws_cons.cell(row=row, column=4, value=item['nome']).border = thin_border
        ws_cons.cell(row=row, column=5, value=item['modulo']).border = thin_border
        ws_cons.cell(row=row, column=6, value=item['complexidade']).border = thin_border
        ws_cons.cell(row=row, column=7, value=item['pf']).border = thin_border
        for col in [1, 2, 3, 6, 7]:
            ws_cons.cell(row=row, column=col).alignment = center_align
        row += 1
        counter += 1

    # SE
    for item in SE_BACKEND:
        ws_cons.cell(row=row, column=1, value=counter).border = thin_border
        ws_cons.cell(row=row, column=2, value="Backend").border = thin_border
        ws_cons.cell(row=row, column=3, value="SE").border = thin_border
        ws_cons.cell(row=row, column=4, value=item['nome']).border = thin_border
        ws_cons.cell(row=row, column=5, value=item['modulo']).border = thin_border
        ws_cons.cell(row=row, column=6, value=item['complexidade']).border = thin_border
        ws_cons.cell(row=row, column=7, value=item['pf']).border = thin_border
        for col in [1, 2, 3, 6, 7]:
            ws_cons.cell(row=row, column=col).alignment = center_align
        row += 1
        counter += 1

    # CE
    for item in CE_BACKEND:
        ws_cons.cell(row=row, column=1, value=counter).border = thin_border
        ws_cons.cell(row=row, column=2, value="Backend").border = thin_border
        ws_cons.cell(row=row, column=3, value="CE").border = thin_border
        ws_cons.cell(row=row, column=4, value=item['nome']).border = thin_border
        ws_cons.cell(row=row, column=5, value=item['modulo']).border = thin_border
        ws_cons.cell(row=row, column=6, value=item['complexidade']).border = thin_border
        ws_cons.cell(row=row, column=7, value=item['pf']).border = thin_border
        for col in [1, 2, 3, 6, 7]:
            ws_cons.cell(row=row, column=col).alignment = center_align
        row += 1
        counter += 1

    # Serviços
    for item in SERVICES_DATA:
        ws_cons.cell(row=row, column=1, value=counter).border = thin_border
        ws_cons.cell(row=row, column=2, value="Backend").border = thin_border
        ws_cons.cell(row=row, column=3, value="Serviço").border = thin_border
        ws_cons.cell(row=row, column=4, value=item['nome']).border = thin_border
        ws_cons.cell(row=row, column=5, value=item['modulo']).border = thin_border
        ws_cons.cell(row=row, column=6, value=item['complexidade']).border = thin_border
        ws_cons.cell(row=row, column=7, value=item['pf']).border = thin_border
        for col in [1, 2, 3, 6, 7]:
            ws_cons.cell(row=row, column=col).alignment = center_align
        row += 1
        counter += 1

    # Frontend
    for item in FRONTEND_DATA:
        ws_cons.cell(row=row, column=1, value=counter).border = thin_border
        ws_cons.cell(row=row, column=2, value="Frontend").border = thin_border
        ws_cons.cell(row=row, column=3, value=item['tipo']).border = thin_border
        ws_cons.cell(row=row, column=4, value=item['nome']).border = thin_border
        ws_cons.cell(row=row, column=5, value=item['modulo']).border = thin_border
        ws_cons.cell(row=row, column=6, value=item['complexidade']).border = thin_border
        ws_cons.cell(row=row, column=7, value=item['pf']).border = thin_border
        for col in [1, 2, 3, 6, 7]:
            ws_cons.cell(row=row, column=col).alignment = center_align
        row += 1
        counter += 1

    # Infraestrutura
    for item in INFRA_DATA:
        ws_cons.cell(row=row, column=1, value=counter).border = thin_border
        ws_cons.cell(row=row, column=2, value="Infra").border = thin_border
        ws_cons.cell(row=row, column=3, value=item['tipo']).border = thin_border
        ws_cons.cell(row=row, column=4, value=item['nome']).border = thin_border
        ws_cons.cell(row=row, column=5, value=item['modulo']).border = thin_border
        ws_cons.cell(row=row, column=6, value=item['complexidade']).border = thin_border
        ws_cons.cell(row=row, column=7, value=item['pf']).border = thin_border
        for col in [1, 2, 3, 6, 7]:
            ws_cons.cell(row=row, column=col).alignment = center_align
        row += 1
        counter += 1

    # Linha zebrada
    for r in range(2, row):
        if r % 2 == 0:
            for col in range(1, 8):
                ws_cons.cell(row=r, column=col).fill = alt_fill

    # Total
    ws_cons.cell(row=row, column=6, value="TOTAL").font = Font(bold=True)
    ws_cons.cell(row=row, column=7, value=pf_total).font = Font(bold=True)
    for col in range(1, 8):
        ws_cons.cell(row=row, column=col).fill = total_fill
        ws_cons.cell(row=row, column=col).border = thin_border

    ws_cons.column_dimensions['A'].width = 5
    ws_cons.column_dimensions['B'].width = 12
    ws_cons.column_dimensions['C'].width = 12
    ws_cons.column_dimensions['D'].width = 55
    ws_cons.column_dimensions['E'].width = 15
    ws_cons.column_dimensions['F'].width = 14
    ws_cons.column_dimensions['G'].width = 8

    # Salvar
    wb.save('C:/Projeto_reavaliacao/APF_Detalhamento_Atividades.xlsx')
    print("Planilha Excel gerada: APF_Detalhamento_Atividades.xlsx")


if __name__ == "__main__":
    print("Gerando documentos de Análise de Pontos de Função...")
    print("-" * 50)
    create_word_document()
    create_excel_spreadsheet()
    print("-" * 50)
    print("Documentos gerados com sucesso!")
    print("\nArquivos criados:")
    print("  1. APF_Sistema_Cotacao.docx (Documento Word)")
    print("  2. APF_Detalhamento_Atividades.xlsx (Planilha Excel)")
