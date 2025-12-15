#!/usr/bin/env python3
"""
Script para gerar documentação Word do fluxo de cotação
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def add_heading(doc, text, level=1):
    """Adiciona título com formatação"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph(doc, text, bold=False, italic=False, color=None):
    """Adiciona parágrafo com formatação"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)

    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color:
        run.font.color.rgb = color

    return p

def add_step_box(doc, step_number, title, description, details):
    """Adiciona caixa de etapa formatada"""
    # Título da etapa
    p = doc.add_paragraph()
    run = p.add_run(f"PASSO {step_number}: {title}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 51, 102)  # Azul escuro
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)

    # Descrição resumida
    p = doc.add_paragraph()
    run = p.add_run(description)
    run.font.size = Pt(11)
    run.italic = True
    run.font.color.rgb = RGBColor(68, 68, 68)  # Cinza escuro
    p.paragraph_format.space_after = Pt(6)

    # Detalhes
    for detail in details:
        p = doc.add_paragraph(detail, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(3)

    # Linha separadora
    doc.add_paragraph("_" * 80)

def create_documentation():
    """Cria o documento Word com a documentação"""
    doc = Document()

    # ============= CAPA =============
    # Título
    title = doc.add_heading('SISTEMA DE COTAÇÃO AUTOMATIZADA', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtítulo
    subtitle = doc.add_heading('Documentação Técnica Completa', 2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    # Informações do documento
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Light Grid Accent 1'

    info_data = [
        ('Data', '12 de Dezembro de 2025'),
        ('Versão', '1.0'),
        ('Sistema', 'Plataforma de Cotação Automatizada'),
        ('Linguagem', 'Python + TypeScript'),
        ('Integrações', 'Claude AI, SerpAPI, Playwright')
    ]

    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value

    doc.add_page_break()

    # ============= SUMÁRIO =============
    add_heading(doc, 'Índice', 1)

    toc_items = [
        '1. Introdução',
        '2. Visão Geral do Sistema',
        '3. Fluxo Completo de Cotação (Passo a Passo)',
        '4. Fluxo com Imagem',
        '5. Fluxo com Texto',
        '6. Tecnologias Utilizadas',
        '7. Custos e Integrações',
        '8. Tratamento de Erros',
        '9. Otimizações Implementadas'
    ]

    for item in toc_items:
        doc.add_paragraph(item, style='List Number')

    doc.add_page_break()

    # ============= INTRODUÇÃO =============
    add_heading(doc, '1. Introdução', 1)

    add_paragraph(doc,
        'Este documento descreve de forma detalhada o funcionamento do Sistema de '
        'Cotação Automatizada, uma plataforma que permite aos usuários obterem '
        'cotações de preços de produtos de forma rápida e automatizada.')

    add_paragraph(doc, '')

    add_paragraph(doc,
        'O sistema aceita dois tipos de entrada:', bold=True)

    doc.add_paragraph('Imagem do produto (foto tirada pelo usuário)', style='List Bullet')
    doc.add_paragraph('Descrição textual do produto', style='List Bullet')

    add_paragraph(doc, '')

    add_paragraph(doc,
        'A partir dessa entrada, o sistema realiza todo o processo de forma automática, '
        'desde a identificação do produto até a coleta de preços em múltiplas lojas online, '
        'retornando ao usuário uma análise completa com preços médios, mínimos e máximos.')

    doc.add_page_break()

    # ============= VISÃO GERAL =============
    add_heading(doc, '2. Visão Geral do Sistema', 1)

    add_heading(doc, '2.1 Arquitetura', 2)

    add_paragraph(doc,
        'O sistema é composto por três camadas principais:')

    doc.add_paragraph('Frontend: Interface web construída com Next.js e React', style='List Bullet')
    doc.add_paragraph('Backend API: Servidor FastAPI em Python que gerencia requisições', style='List Bullet')
    doc.add_paragraph('Worker Assíncrono: Celery worker que processa cotações em background', style='List Bullet')

    add_paragraph(doc, '')

    add_heading(doc, '2.2 Bancos de Dados e Filas', 2)

    doc.add_paragraph('PostgreSQL: Armazena todas as cotações, preços e dados estruturados', style='List Bullet')
    doc.add_paragraph('Redis: Gerencia a fila de tarefas assíncronas', style='List Bullet')

    add_paragraph(doc, '')

    add_heading(doc, '2.3 Integrações Externas', 2)

    doc.add_paragraph('Anthropic Claude API: Inteligência artificial para análise de produtos', style='List Bullet')
    doc.add_paragraph('SerpAPI: Busca de produtos no Google Shopping', style='List Bullet')
    doc.add_paragraph('Playwright: Automação de navegador para captura de preços', style='List Bullet')

    doc.add_page_break()

    # ============= FLUXO COMPLETO =============
    add_heading(doc, '3. Fluxo Completo de Cotação (Passo a Passo)', 1)

    add_paragraph(doc,
        'O processo de cotação segue uma sequência bem definida de etapas. '
        'A seguir, descreveremos cada passo em detalhes:', bold=True)

    add_paragraph(doc, '')

    # ========== PASSO 1 ==========
    add_step_box(doc, 1,
        'ENTRADA DO USUÁRIO',
        'O usuário acessa o sistema e fornece informações sobre o produto que deseja cotar.',
        [
            'O usuário pode enviar UMA OU MAIS IMAGENS do produto (fotos tiradas com celular ou câmera)',
            'Ou pode DIGITAR uma descrição textual do produto (ex: "Cadeira gamer preta com apoio de braço")',
            'Opcionalmente, pode informar: código do item, local da cotação, nome do pesquisador',
            'Pode vincular a cotação a um projeto existente no sistema',
            'As imagens enviadas são salvas no servidor com hash SHA256 para garantir integridade'
        ])

    # ========== PASSO 2 ==========
    add_step_box(doc, 2,
        'CRIAÇÃO DA REQUISIÇÃO',
        'O sistema recebe os dados e cria um registro de cotação no banco de dados.',
        [
            'Backend valida que pelo menos imagem OU texto foi fornecido',
            'Cria registro na tabela "quote_requests" com status PROCESSING',
            'Salva as imagens no diretório "storage/input_images/"',
            'Gera um ID único para a cotação (ex: 76)',
            'Retorna o ID da cotação para o frontend',
            'Enfileira a tarefa de processamento no Celery (processamento assíncrono)'
        ])

    # ========== PASSO 3 ==========
    add_step_box(doc, 3,
        'INICIALIZAÇÃO DO PROCESSAMENTO',
        'O Celery Worker pega a tarefa da fila e inicia o processamento.',
        [
            'Progresso atualizado: 5% - "Carregando configurações e integrações..."',
            'Carrega a chave de API da Anthropic (Claude)',
            'Define qual modelo de IA será usado (ex: claude-3-5-sonnet-20241022)',
            'Carrega chave de API do SerpAPI para busca de produtos',
            'Carrega parâmetros configuráveis: número de cotações, tolerância de outliers',
            'Busca no banco de dados se há imagens anexadas à cotação'
        ])

    # ========== PASSO 4 ==========
    add_step_box(doc, 4,
        'ANÁLISE INTELIGENTE COM IA',
        'O sistema envia os dados para a inteligência artificial Claude para análise.',
        [
            'Progresso: 10%',
            'SE ENVIOU IMAGEM: "Processando imagens e extraindo especificações técnicas..."',
            'SE ENVIOU TEXTO: "Analisando descrição e identificando produto..."',
            '',
            'O que o Claude faz:',
            '  • Se houver IMAGEM: Analisa visualmente a foto, identifica o produto, lê textos na imagem',
            '  • Se houver TEXTO: Interpreta a descrição e identifica características',
            '  • Determina o TIPO do produto (ex: "Cadeira Gamer", "Notebook", "Mouse")',
            '  • Extrai ESPECIFICAÇÕES TÉCNICAS (cor, material, tamanho, marca, modelo, etc.)',
            '  • Gera uma QUERY DE BUSCA OTIMIZADA para usar no Google Shopping',
            '  • Gera QUERIES ALTERNATIVAS caso a busca principal não retorne resultados',
            '',
            'Exemplo de resultado:',
            '  Tipo: "Cadeira Gamer"',
            '  Especificações: { "cor": "preta", "material": "couro sintético", "ajuste": "sim" }',
            '  Query principal: "cadeira gamer preta couro sintetico ergonomica"',
            '  Queries alternativas: ["cadeira gamer preta reclinavel", "poltrona gamer preta"]'
        ])

    # ========== PASSO 5 ==========
    add_step_box(doc, 5,
        'REGISTRO DE CUSTOS DA IA',
        'O sistema registra quanto custou a análise do Claude.',
        [
            'Progresso: 30% - "Análise completa - [N] tokens processados pela IA"',
            'Claude retorna quantos "tokens" foram processados (tokens = unidades de texto/imagem)',
            'Sistema calcula o custo baseado no modelo usado:',
            '  • Claude 3.5 Sonnet: $3.00 por milhão de tokens de entrada',
            '  • Claude 3.5 Sonnet: $15.00 por milhão de tokens de saída',
            'Converte o custo de USD para BRL usando taxa configurada',
            'Salva o custo na tabela "financial_transactions" para controle financeiro',
            '',
            'Exemplo: Se processou 1.500 tokens, custou aproximadamente R$ 0,15'
        ])

    # ========== PASSO 6 ==========
    add_step_box(doc, 6,
        'PREPARAÇÃO PARA BUSCA',
        'O sistema se prepara para buscar produtos em marketplaces.',
        [
            'Progresso: 40% - "Preparando busca de preços em lojas online..."',
            'Carrega a chave de API do SerpAPI',
            'Define quantos preços buscar (configurável, padrão: 3 preços)',
            'Define tolerância para outliers (preços muito fora da média)',
            'Configura localização da busca (Brasil)',
            'Prepara a query de busca otimizada que o Claude gerou'
        ])

    # ========== PASSO 7 ==========
    add_step_box(doc, 7,
        'BUSCA DE PRODUTOS NO GOOGLE SHOPPING',
        'O sistema busca o produto no Google Shopping via SerpAPI.',
        [
            'Progresso: 50% - "Buscando \'[query]\' em marketplaces..."',
            '',
            'Estratégia de busca em 2 fases:',
            '',
            'FASE 1 - Google Shopping:',
            '  • Faz 1 chamada à API do Google Shopping',
            '  • Envia a query otimizada (ex: "cadeira gamer preta couro sintetico")',
            '  • Retorna aproximadamente 20 produtos candidatos',
            '  • Cada produto tem: título, preço estimado, link',
            '',
            'FASE 2 - Filtragem Inteligente:',
            '  • Analisa os preços dos 20 produtos',
            '  • Remove produtos com preços muito discrepantes (outliers)',
            '  • Seleciona os TOP 3 produtos mais relevantes',
            '',
            'FASE 3 - Detalhamento (Immersive Product):',
            '  • Para cada um dos 3 produtos selecionados:',
            '  • Faz chamada "Immersive Product" para pegar URL exata da loja',
            '  • Total de chamadas: 1 Shopping + 3 Immersive = 4 chamadas de API',
            '',
            'Otimização: Fazemos apenas 4 chamadas em vez de 21, economizando 80% do custo!'
        ])

    # ========== PASSO 8 ==========
    add_step_box(doc, 8,
        'EXTRAÇÃO DE PREÇOS DAS LOJAS',
        'O sistema acessa cada loja e captura o preço real do produto.',
        [
            'Progresso: 60% - "Acessando [N] lojas e capturando preços..."',
            '',
            'Para cada uma das 3 lojas selecionadas:',
            '',
            '1. ABERTURA DA PÁGINA:',
            '   • Usa Playwright (navegador Chromium automatizado)',
            '   • Abre a URL da loja em modo headless (sem interface gráfica)',
            '   • Aguarda a página carregar completamente',
            '',
            '2. CAPTURA DE SCREENSHOT:',
            '   • Tira uma foto (screenshot) da página completa',
            '   • Salva a imagem no servidor (storage/screenshots/)',
            '   • Isso serve como prova/evidência do preço',
            '',
            '3. EXTRAÇÃO DO PREÇO:',
            '   • Tenta encontrar o preço usando múltiplos métodos:',
            '     a) JSON-LD schema (dados estruturados da página)',
            '     b) Meta tags OpenGraph (tags de compartilhamento social)',
            '     c) Seletores CSS específicos (.price, .product-price, etc.)',
            '     d) Expressões regulares no HTML (busca por R$ XX,XX)',
            '   • Pega o PRIMEIRO método que funcionar',
            '   • Valida que o preço é maior que R$ 1,00',
            '',
            '4. SALVAMENTO:',
            '   • Salva na tabela "quote_sources":',
            '     - URL da loja',
            '     - Preço encontrado',
            '     - Título da página',
            '     - ID do screenshot',
            '     - Método de extração usado',
            '     - Data/hora da captura',
            '',
            'Paralelização: Processa até 3 lojas simultâneas para acelerar'
        ])

    # ========== PASSO 9 ==========
    add_step_box(doc, 9,
        'ANÁLISE ESTATÍSTICA DOS PREÇOS',
        'O sistema analisa os preços coletados e calcula estatísticas.',
        [
            'Progresso: 80% - "Analisando [N] preços coletados e calculando média..."',
            '',
            '1. DETECÇÃO DE OUTLIERS:',
            '   • Usa método IQR (Interquartile Range)',
            '   • Calcula quartis Q1 (25%) e Q3 (75%)',
            '   • Calcula IQR = Q3 - Q1',
            '   • Define limites: inferior = Q1 - 1.5*IQR, superior = Q3 + 1.5*IQR',
            '   • Marca preços fora desses limites como "outliers"',
            '   • Outliers são ignorados no cálculo final',
            '',
            '2. CÁLCULO DE ESTATÍSTICAS:',
            '   • Preço MÉDIO: soma dos preços aceitos ÷ quantidade',
            '   • Preço MÍNIMO: menor preço encontrado',
            '   • Preço MÁXIMO: maior preço encontrado',
            '',
            '3. REGISTRO DE CUSTOS:',
            '   • Registra custo da busca no SerpAPI (~R$ 0,50 por cotação)',
            '   • Salva em "financial_transactions" para controle',
            '',
            'Exemplo:',
            '  Preços encontrados: R$ 369,90 | R$ 379,05 | R$ 599,00 (outlier)',
            '  Resultado: Média = R$ 374,48 | Mín = R$ 369,90 | Máx = R$ 379,05'
        ])

    # ========== PASSO 10 ==========
    add_step_box(doc, 10,
        'FINALIZAÇÃO DA COTAÇÃO',
        'O sistema salva todos os resultados e marca a cotação como concluída.',
        [
            'Progresso: 95% - "Salvando resultados e finalizando cotação..."',
            '',
            'Operações finais:',
            '  • Salva todas as estatísticas no banco de dados',
            '  • Marca status da cotação como DONE (concluída)',
            '  • Define progresso como 100%',
            '  • Registra data/hora de conclusão',
            '  • Libera recursos (fecha navegador, limpa memória)',
            '',
            'Progresso: 100% - "Cotação concluída! Preços capturados e analisados com sucesso."',
            '',
            'A partir deste momento:',
            '  • Usuário pode visualizar os resultados',
            '  • Pode gerar PDF da cotação (sob demanda)',
            '  • Pode vincular com tabela de materiais do projeto',
            '  • Pode fazer nova cotação baseada nesta'
        ])

    doc.add_page_break()

    # ============= FLUXO COM IMAGEM =============
    add_heading(doc, '4. Fluxo Detalhado: Cotação por Imagem', 1)

    add_paragraph(doc,
        'Quando o usuário envia uma IMAGEM, o fluxo tem características específicas:')

    add_paragraph(doc, '')

    add_heading(doc, 'Exemplo Prático: Foto de uma Cadeira Gamer', 2)

    add_paragraph(doc, 'ENTRADA:', bold=True)
    doc.add_paragraph('Usuário tira foto de uma cadeira gamer com o celular', style='List Bullet')
    doc.add_paragraph('Imagem mostra: cadeira preta, com logo "X", apoio de braço regulável', style='List Bullet')
    doc.add_paragraph('Pode ter texto visível na foto (marca, modelo)', style='List Bullet')

    add_paragraph(doc, '')
    add_paragraph(doc, 'O QUE O CLAUDE FAZ:', bold=True)

    doc.add_paragraph('Analisa visualmente a imagem usando visão computacional', style='List Number')
    doc.add_paragraph('Identifica que é uma cadeira gamer (tipo de produto)', style='List Number')
    doc.add_paragraph('Detecta características visuais: cor preta, design gaming, estrutura', style='List Number')
    doc.add_paragraph('Lê qualquer texto visível na imagem (OCR automático)', style='List Number')
    doc.add_paragraph('Extrai especificações técnicas visíveis', style='List Number')
    doc.add_paragraph('Gera query de busca: "cadeira gamer preta ergonomica reclinavel"', style='List Number')

    add_paragraph(doc, '')
    add_paragraph(doc, 'RESULTADO DA ANÁLISE:', bold=True)

    resultado_table = doc.add_table(rows=5, cols=2)
    resultado_table.style = 'Light Grid Accent 1'

    resultado_data = [
        ('Tipo', 'Cadeira Gamer'),
        ('Cor', 'Preta'),
        ('Material', 'Couro sintético (detectado visualmente)'),
        ('Características', 'Ergonômica, reclinável, apoio regulável'),
        ('Query gerada', 'cadeira gamer preta ergonomica reclinavel')
    ]

    for i, (label, value) in enumerate(resultado_data):
        resultado_table.rows[i].cells[0].text = label
        resultado_table.rows[i].cells[1].text = value

    add_paragraph(doc, '')
    add_paragraph(doc,
        'Com essa query otimizada, o sistema segue para a busca no Google Shopping '
        'e coleta os preços conforme descrito nos passos anteriores.')

    doc.add_page_break()

    # ============= FLUXO COM TEXTO =============
    add_heading(doc, '5. Fluxo Detalhado: Cotação por Texto', 1)

    add_paragraph(doc,
        'Quando o usuário DIGITA uma descrição, o fluxo é otimizado para análise textual:')

    add_paragraph(doc, '')

    add_heading(doc, 'Exemplo Prático: Descrição Digitada', 2)

    add_paragraph(doc, 'ENTRADA:', bold=True)
    texto_input = doc.add_paragraph()
    texto_input.add_run('Usuário digita: ').bold = True
    texto_input.add_run('"Notebook Dell Inspiron i5 11ª geração, 16GB RAM, SSD 512GB, tela 15.6 polegadas"')

    add_paragraph(doc, '')
    add_paragraph(doc, 'O QUE O CLAUDE FAZ:', bold=True)

    doc.add_paragraph('Interpreta o texto usando processamento de linguagem natural (NLP)', style='List Number')
    doc.add_paragraph('Identifica tipo de produto: Notebook', style='List Number')
    doc.add_paragraph('Extrai especificações mencionadas:', style='List Number')

    specs = doc.add_paragraph(style='List Bullet 2')
    specs.add_run('  • Marca: Dell\n')
    specs.add_run('  • Linha: Inspiron\n')
    specs.add_run('  • Processador: Intel i5 11ª geração\n')
    specs.add_run('  • RAM: 16GB\n')
    specs.add_run('  • Armazenamento: SSD 512GB\n')
    specs.add_run('  • Tela: 15.6"')

    doc.add_paragraph('Normaliza termos técnicos (ex: "i5" → "Intel Core i5")', style='List Number')
    doc.add_paragraph('Gera query otimizada removendo palavras desnecessárias', style='List Number')
    doc.add_paragraph('Cria queries alternativas para garantir resultados', style='List Number')

    add_paragraph(doc, '')
    add_paragraph(doc, 'RESULTADO DA ANÁLISE:', bold=True)

    resultado_texto_table = doc.add_table(rows=4, cols=2)
    resultado_texto_table.style = 'Light Grid Accent 1'

    resultado_texto_data = [
        ('Tipo', 'Notebook'),
        ('Marca/Modelo', 'Dell Inspiron'),
        ('Especificações principais', 'i5 11ª gen, 16GB RAM, SSD 512GB, 15.6"'),
        ('Query otimizada', 'notebook dell inspiron i5 11 geracao 16gb ssd 512gb')
    ]

    for i, (label, value) in enumerate(resultado_texto_data):
        resultado_texto_table.rows[i].cells[0].text = label
        resultado_texto_table.rows[i].cells[1].text = value

    add_paragraph(doc, '')
    add_paragraph(doc,
        'IMPORTANTE: O Claude é treinado para entender português brasileiro e termos '
        'técnicos específicos, reconhecendo abreviações como "RAM", "SSD", "i5", etc.')

    doc.add_page_break()

    # ============= TECNOLOGIAS =============
    add_heading(doc, '6. Tecnologias Utilizadas', 1)

    add_heading(doc, '6.1 Anthropic Claude (Inteligência Artificial)', 2)

    add_paragraph(doc, 'O que é:', bold=True)
    add_paragraph(doc,
        'Claude é uma IA avançada da Anthropic, capaz de processar texto e imagens, '
        'compreender contexto e extrair informações estruturadas.')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Por que usamos:', bold=True)
    doc.add_paragraph('Análise precisa de produtos a partir de fotos ou descrições', style='List Bullet')
    doc.add_paragraph('Extração automática de especificações técnicas', style='List Bullet')
    doc.add_paragraph('Geração de queries de busca otimizadas', style='List Bullet')
    doc.add_paragraph('Compreensão de termos técnicos e contexto brasileiro', style='List Bullet')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Modelo usado:', bold=True)
    add_paragraph(doc, 'claude-3-5-sonnet-20241022 (configurável)')

    add_paragraph(doc, '')

    add_heading(doc, '6.2 SerpAPI (Google Shopping)', 2)

    add_paragraph(doc, 'O que é:', bold=True)
    add_paragraph(doc,
        'SerpAPI é um serviço que fornece acesso programático aos resultados do '
        'Google Shopping, retornando dados estruturados em JSON.')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Por que usamos:', bold=True)
    doc.add_paragraph('Busca de produtos em múltiplas lojas simultaneamente', style='List Bullet')
    doc.add_paragraph('Retorna preços, links e informações de produtos', style='List Bullet')
    doc.add_paragraph('Mais confiável que fazer scraping direto do Google', style='List Bullet')
    doc.add_paragraph('API oficial com alta disponibilidade', style='List Bullet')

    add_paragraph(doc, '')

    add_heading(doc, '6.3 Playwright (Web Scraping)', 2)

    add_paragraph(doc, 'O que é:', bold=True)
    add_paragraph(doc,
        'Playwright é uma ferramenta de automação de navegadores que permite '
        'controlar o Chrome/Firefox/Safari de forma programática.')

    add_paragraph(doc, '')
    add_paragraph(doc, 'Por que usamos:', bold=True)
    doc.add_paragraph('Acessa as páginas das lojas como um usuário real', style='List Bullet')
    doc.add_paragraph('Captura screenshots das páginas para evidência', style='List Bullet')
    doc.add_paragraph('Extrai preços mesmo de sites com JavaScript dinâmico', style='List Bullet')
    doc.add_paragraph('Suporta múltiplos métodos de extração de preço', style='List Bullet')

    add_paragraph(doc, '')

    add_heading(doc, '6.4 PostgreSQL (Banco de Dados)', 2)

    add_paragraph(doc, 'Armazena todos os dados:', bold=True)
    doc.add_paragraph('Cotações e seus status', style='List Bullet')
    doc.add_paragraph('Preços coletados de cada loja', style='List Bullet')
    doc.add_paragraph('Screenshots e arquivos', style='List Bullet')
    doc.add_paragraph('Transações financeiras (custos)', style='List Bullet')
    doc.add_paragraph('Projetos e vinculações', style='List Bullet')

    add_paragraph(doc, '')

    add_heading(doc, '6.5 Redis + Celery (Fila de Tarefas)', 2)

    add_paragraph(doc, 'Por que processamento assíncrono:', bold=True)
    add_paragraph(doc,
        'Como a cotação pode demorar 20-30 segundos, o processamento é feito em '
        'background para não bloquear o usuário. Ele pode continuar navegando '
        'enquanto a cotação processa.')

    doc.add_paragraph('Redis: Gerencia a fila de tarefas pendentes', style='List Bullet')
    doc.add_paragraph('Celery: Worker que executa as tarefas em background', style='List Bullet')

    doc.add_page_break()

    # ============= CUSTOS =============
    add_heading(doc, '7. Custos e Integrações', 1)

    add_heading(doc, '7.1 Custo por Cotação', 2)

    custos_table = doc.add_table(rows=4, cols=3)
    custos_table.style = 'Light Grid Accent 1'

    # Header
    custos_table.rows[0].cells[0].text = 'Serviço'
    custos_table.rows[0].cells[1].text = 'Custo Aproximado'
    custos_table.rows[0].cells[2].text = 'Observação'

    custos_data = [
        ('Anthropic Claude', 'R$ 0,10 - R$ 0,30', 'Varia com tamanho da imagem/texto'),
        ('SerpAPI', 'R$ 0,50', 'Fixo por cotação (4 chamadas)'),
        ('TOTAL', 'R$ 0,60 - R$ 0,80', 'Por cotação completa')
    ]

    for i, (servico, custo, obs) in enumerate(custos_data, 1):
        custos_table.rows[i].cells[0].text = servico
        custos_table.rows[i].cells[1].text = custo
        custos_table.rows[i].cells[2].text = obs

    add_paragraph(doc, '')
    add_paragraph(doc,
        'IMPORTANTE: Todos os custos são registrados na tabela "financial_transactions" '
        'para controle financeiro detalhado.')

    add_paragraph(doc, '')

    add_heading(doc, '7.2 Otimizações de Custo', 2)

    add_paragraph(doc, 'Redução de 80% em chamadas de API:', bold=True)

    otim_table = doc.add_table(rows=3, cols=3)
    otim_table.style = 'Medium Grid 3 Accent 1'

    otim_table.rows[0].cells[0].text = 'Estratégia'
    otim_table.rows[0].cells[1].text = 'Antes'
    otim_table.rows[0].cells[2].text = 'Depois'

    otim_table.rows[1].cells[0].text = 'Chamadas SerpAPI'
    otim_table.rows[1].cells[1].text = '1 + 20 = 21 calls'
    otim_table.rows[1].cells[2].text = '1 + 3 = 4 calls'

    otim_table.rows[2].cells[0].text = 'Custo por cotação'
    otim_table.rows[2].cells[1].text = 'R$ 2,50'
    otim_table.rows[2].cells[2].text = 'R$ 0,50'

    add_paragraph(doc, '')
    add_paragraph(doc,
        'Como conseguimos: Aplicamos filtro de outliers ANTES de fazer chamadas '
        'Immersive, reduzindo de 20 para apenas 3 produtos.')

    doc.add_page_break()

    # ============= ERROS =============
    add_heading(doc, '8. Tratamento de Erros', 1)

    add_heading(doc, '8.1 Tipos de Erro', 2)

    erros_table = doc.add_table(rows=5, cols=3)
    erros_table.style = 'Light Grid Accent 1'

    erros_table.rows[0].cells[0].text = 'Erro'
    erros_table.rows[0].cells[1].text = 'Causa'
    erros_table.rows[0].cells[2].text = 'Ação do Sistema'

    erros_data = [
        ('Claude API falhou', 'Timeout, rate limit, API fora do ar', 'Retry automático 3x, depois marca como ERROR'),
        ('Nenhum preço encontrado', 'Produto muito específico ou raro', 'Marca como ERROR com mensagem clara'),
        ('SerpAPI sem quota', 'Limite mensal atingido', 'Marca como ERROR, notifica admin'),
        ('Página não carregou', 'Site lento ou fora do ar', 'Pula essa loja, tenta próxima')
    ]

    for i, (erro, causa, acao) in enumerate(erros_data, 1):
        erros_table.rows[i].cells[0].text = erro
        erros_table.rows[i].cells[1].text = causa
        erros_table.rows[i].cells[2].text = acao

    add_paragraph(doc, '')

    add_heading(doc, '8.2 Status da Cotação', 2)

    status_items = [
        'PROCESSING: Cotação em andamento',
        'DONE: Cotação concluída com sucesso',
        'ERROR: Ocorreu erro durante processamento',
        'CANCELLED: Usuário cancelou a cotação'
    ]

    for item in status_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ============= OTIMIZAÇÕES =============
    add_heading(doc, '9. Otimizações Implementadas', 1)

    add_heading(doc, '9.1 Performance', 2)

    doc.add_paragraph('Paralelização de scraping: 3 lojas simultâneas (3x mais rápido)', style='List Bullet')
    doc.add_paragraph('Cache Redis: Queries repetidas retornam resultados em < 1s', style='List Bullet')
    doc.add_paragraph('Database indexes: Buscas otimizadas com índices em campos chave', style='List Bullet')
    doc.add_paragraph('Lazy loading: Screenshots carregados sob demanda', style='List Bullet')

    add_paragraph(doc, '')

    add_heading(doc, '9.2 Tempo Médio de Processamento', 2)

    tempo_table = doc.add_table(rows=9, cols=2)
    tempo_table.style = 'Light Grid Accent 1'

    tempo_table.rows[0].cells[0].text = 'Etapa'
    tempo_table.rows[0].cells[1].text = 'Tempo'

    tempo_data = [
        ('Inicialização', '1-2s'),
        ('Análise Claude (IA)', '3-5s'),
        ('Preparação busca', '0.5s'),
        ('Busca SerpAPI', '2-3s'),
        ('Extração preços (3 lojas)', '10-15s'),
        ('Cálculo estatísticas', '0.5s'),
        ('Finalização', '0.5s'),
        ('TOTAL', '18-27 segundos')
    ]

    for i, (etapa, tempo) in enumerate(tempo_data, 1):
        tempo_table.rows[i].cells[0].text = etapa
        tempo_table.rows[i].cells[1].text = tempo

    add_paragraph(doc, '')

    add_heading(doc, '9.3 Confiabilidade', 2)

    doc.add_paragraph('Retry automático em falhas de API externa (3 tentativas)', style='List Bullet')
    doc.add_paragraph('Screenshots como prova dos preços coletados', style='List Bullet')
    doc.add_paragraph('Detecção de outliers para remover preços suspeitos', style='List Bullet')
    doc.add_paragraph('Validação de preços (mínimo R$ 1,00)', style='List Bullet')
    doc.add_paragraph('Registro completo de custos para auditoria', style='List Bullet')

    doc.add_page_break()

    # ============= CONCLUSÃO =============
    add_heading(doc, '10. Conclusão', 1)

    add_paragraph(doc,
        'O Sistema de Cotação Automatizada representa uma solução moderna e eficiente '
        'para o processo de pesquisa de preços. Através da integração de tecnologias '
        'de ponta como:', bold=True)

    add_paragraph(doc, '')

    doc.add_paragraph('Inteligência Artificial (Claude) para análise de produtos', style='List Bullet')
    doc.add_paragraph('APIs especializadas (SerpAPI) para busca em marketplaces', style='List Bullet')
    doc.add_paragraph('Automação de navegadores (Playwright) para coleta precisa de preços', style='List Bullet')
    doc.add_paragraph('Processamento assíncrono (Celery) para experiência fluida do usuário', style='List Bullet')

    add_paragraph(doc, '')

    add_paragraph(doc,
        'O sistema consegue processar cotações em menos de 30 segundos, com custos '
        'controlados (< R$ 1,00 por cotação) e alta precisão nos resultados.')

    add_paragraph(doc, '')

    add_paragraph(doc, 'Principais diferenciais:', bold=True)

    doc.add_paragraph('Aceita tanto IMAGENS quanto TEXTO como entrada', style='List Number')
    doc.add_paragraph('Análise inteligente usando IA de última geração', style='List Number')
    doc.add_paragraph('Coleta preços de múltiplas lojas automaticamente', style='List Number')
    doc.add_paragraph('Fornece evidência visual (screenshots) dos preços', style='List Number')
    doc.add_paragraph('Calcula estatísticas e remove outliers automaticamente', style='List Number')
    doc.add_paragraph('Registra todos os custos para controle financeiro', style='List Number')
    doc.add_paragraph('Interface em tempo real mostrando progresso da cotação', style='List Number')

    add_paragraph(doc, '')
    add_paragraph(doc, '')

    final = doc.add_paragraph()
    run = final.add_run('Este documento descreve a versão atual do sistema (v1.0). ')
    run = final.add_run('Para dúvidas técnicas ou sugestões de melhorias, consulte o time de desenvolvimento.')

    # Salvar documento
    output_path = 'C:\\Projeto_reavaliacao\\DOCUMENTACAO_COMPLETA_COTACAO.docx'
    doc.save(output_path)
    print(f"Documento criado com sucesso: {output_path}")
    return output_path

if __name__ == '__main__':
    create_documentation()
